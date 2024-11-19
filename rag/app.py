import streamlit as st
import openai
import os
from dotenv import load_dotenv
from langchain_community.vectorstores import FAISS
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.llms import OpenAI
from pymongo import MongoClient
from datetime import datetime
import uuid
import hashlib
from sklearn.preprocessing import StandardScaler
import numpy as np
from docx import Document
import time
from langsmith import Client
from typing import List, Dict, Optional
import json
import sys
sys.path.append("../stt")
from gcp import record_audio as gcp_record
from gcp import audio_to_text as gcp_transcribe
from whisper_stt import AudioTranscriber
import extra_streamlit_components as stx
from datetime import timedelta

class Session:
    def __init__(self, session_id, timestamp, consultation_text="", generated_content="", rating=None, feedback=""):
        self.session_id = session_id
        self.timestamp = timestamp
        self.consultation_text = consultation_text
        self.generated_content = generated_content
        self.rating = rating
        self.feedback = feedback

class UserManager:
    def __init__(self, mongo_db):
        self.mongo_db = mongo_db
        if 'user' not in st.session_state:
            st.session_state.user = None
        
        # 쿠키 매니저 초기화
        self.cookie_manager = stx.CookieManager()
        st.write("쿠키 매니저 초기화됨")  # Debug log
        
        # 쿠키에서 세션 복원
        self._restore_session_from_cookie()

    def _restore_session_from_cookie(self):
        """쿠키에서 사용자 세션을 복원합니다."""
        if not st.session_state.user:  # 세션에 사용자가 없을 때만 쿠키 확인
            username = self.cookie_manager.get("username")
            st.write(f"쿠키에서 읽은 username: {username}")  # Debug log
            if username:
                user = self.mongo_db.db.users.find_one({"username": username})
                if user:
                    st.session_state.user = {
                        "username": user["username"],
                        "user_id": str(user["_id"])
                    }
                    st.write("쿠키에서 세션 복원 성공")  # Debug log

    def login(self, username, password):
        # 실제 구현에서는 비밀번호 해싱 필요
        user = self.mongo_db.db.users.find_one({
            "username": username,
            "password": password
        })
        if user:
            st.session_state.user = {
                "username": user["username"],
                "user_id": str(user["_id"])
            }
            # 24시간 유효한 쿠키 설정
            try:
                self.cookie_manager.set(
                    "username",
                    user["username"],
                    expires_at=datetime.now() + timedelta(days=1)
                )
                st.write("쿠키 설정 성공")  # Debug log
                
                # 쿠키가 제대로 설정되었는지 즉시 확인
                saved_username = self.cookie_manager.get("username")
                st.write(f"저장된 쿠키 확인: {saved_username}")  # Debug log
            except Exception as e:
                st.write(f"쿠키 설정 실패: {str(e)}")  # Debug log
            return True
        return False

    def logout(self):
        st.session_state.user = None
        st.session_state.sessions = []
        st.session_state.current_session_id = None
        # 쿠키 삭제
        try:
            self.cookie_manager.delete("username")
            st.write("쿠키 삭제 성공")  # Debug log
        except Exception as e:
            st.write(f"쿠키 삭제 실패: {str(e)}")  # Debug log

    def is_logged_in(self):
        if not st.session_state.user:
            self._restore_session_from_cookie()
        return st.session_state.user is not None

    def get_current_user(self):
        if not st.session_state.user:
            self._restore_session_from_cookie()
        return st.session_state.user

    def register(self, username, password):
        # 이미 존재하는 사용자인지 확인
        existing_user = self.mongo_db.db.users.find_one({"username": username})
        if existing_user:
            return False, "이미 존재하는 사용자입니다."
        
        # 새 사용자 등록
        user = {
            "username": username,
            "password": password,  # 실제 구현시 해시 처리 필요
            "created_at": datetime.now()
        }
        try:
            self.mongo_db.db.users.insert_one(user)
            return True, "사용자 등록이 완료되었습니다."
        except Exception as e:
            return False, f"사용자 등록 중 오류가 발생했습니다: {str(e)}"

class SessionManager:
    def __init__(self):
        if 'sessions' not in st.session_state:
            st.session_state.sessions = []
        if 'current_session_id' not in st.session_state:
            st.session_state.current_session_id = None
        self.db = get_database()
        self.sessions_collection = self.db['sessions']

    def save_session(self, consultation_text, generated_content):
        user = st.session_state.get('user')
        if not user:
            raise ValueError("User not logged in")

        session = Session(
            session_id=str(uuid.uuid4()),
            timestamp=datetime.now(),
            consultation_text=consultation_text,
            generated_content=generated_content
        )
        session.user_id = user['user_id']
        session.username = user['username']
        
        # MongoDB에 저장
        session_data = {
            'session_id': session.session_id,
            'user_id': session.user_id,
            'username': session.username,
            'timestamp': session.timestamp,
            'consultation_text': session.consultation_text,
            'generated_content': session.generated_content,
            'rating': None,
            'feedback': ''
        }
        self.sessions_collection.insert_one(session_data)
        
        st.session_state.sessions.append(session)
        st.session_state.current_session_id = session.session_id
        return session.session_id

    def get_sessions(self):
        """현재 로그인한 사용자의 세션을 MongoDB에서 가져옵니다."""
        user = st.session_state.get('user')
        if not user:
            return []
        
        # MongoDB에서 세션 로드
        session_docs = self.sessions_collection.find(
            {'user_id': user['user_id']},
            sort=[('timestamp', -1)]  # 최신 순으로 정렬
        )
        
        sessions = []
        for doc in session_docs:
            session = Session(
                session_id=doc['session_id'],
                timestamp=doc['timestamp'],
                consultation_text=doc['consultation_text'],
                generated_content=doc['generated_content']
            )
            session.user_id = doc['user_id']
            session.username = doc['username']
            session.rating = doc.get('rating')
            session.feedback = doc.get('feedback', '')
            sessions.append(session)
        
        # 세션 상태 업데이트
        st.session_state.sessions = sessions
        return sessions

    def update_current_session(self, rating=None, feedback=None):
        if not st.session_state.current_session_id:
            return False
        
        update_data = {}
        if rating is not None:
            update_data['rating'] = rating
        if feedback is not None:
            update_data['feedback'] = feedback
        
        if update_data:
            # MongoDB 업데이트
            self.sessions_collection.update_one(
                {'session_id': st.session_state.current_session_id},
                {'$set': update_data}
            )
            
            # 메모리 상의 세션도 업데이트
            for session in st.session_state.sessions:
                if session.session_id == st.session_state.current_session_id:
                    if rating is not None:
                        session.rating = rating
                    if feedback is not None:
                        session.feedback = feedback
                    return True
        return False

def display_sessions():
    """세션 목록을 표시합니다."""
    if not st.session_state.user:
        return

    st.write("### 💬 상담 기록")
    
    # 현재 사용자 정보와 로그아웃 버튼
    col1, col2 = st.columns([2, 1])
    with col1:
        st.markdown(f"**👤 현재 사용자:** {st.session_state.user['username']}")
    with col2:
        if st.button("로그아웃", use_container_width=True):
            st.session_state.generator.user_manager.logout()
            st.rerun()
    
    st.markdown("---")  # 구분선 추가
    
    if st.button("✨ 새로운 상담 시작하기", type="primary", use_container_width=True):
        st.session_state.current_session_id = None
        st.session_state.consultation_text = ""
        st.session_state.generated_complaint = ""
        st.rerun()

    # 세션 목록 가져오기
    sessions = st.session_state.generator.session_manager.get_sessions()
    
    if not sessions:
        st.info("아직 상담 기록이 없습니다.")
        return

    # 세션 목록 표시
    for idx, session in enumerate(reversed(sessions)):
        with st.expander(f"상담 {idx + 1} - {session.timestamp.strftime('%Y-%m-%d %H:%M')}"):
            st.markdown("#### 상담 내용")
            st.write(session.consultation_text)
            
            st.markdown("#### 생성된 소장")
            st.write(session.generated_content)
            
            # 평가 섹션
            st.markdown("#### 평가")
            if hasattr(session, 'rating') and session.rating:
                st.write(f"⭐ 평점: {session.rating}")
            if hasattr(session, 'feedback') and session.feedback:
                st.write(f"💭 피드백: {session.feedback}")
            
            # 현재 세션이 아닌 경우에만 '이 버전 사용' 버튼 표시
            if session.session_id != st.session_state.current_session_id:
                if st.button("이 버전 사용", key=f"use_version_{session.session_id}", use_container_width=True):
                    st.session_state.current_session_id = session.session_id
                    st.rerun()

class STTManager:
    def __init__(self):
        print("Initializing STT Manager...")
        self.whisper_transcriber = AudioTranscriber()
        self.recording = False
        self.last_transcription_time = time.time()
        self.transcription_buffer = []
        
    def start_recording(self, engine="whisper"):
        """녹음을 시작합니다."""
        print(f"Starting recording with {engine} engine...")
        if not self.recording:
            try:
                if engine == "whisper":
                    self.recording = True
                    self.transcription_buffer = []
                    self.last_transcription_time = time.time()
                    self.whisper_transcriber.start_streaming(
                        callback=self.on_transcription_complete
                    )
                    print("Recording started successfully")
                    return True
                else:
                    print("Selected engine not supported yet")
                    return False
            except Exception as e:
                print(f"Error starting recording: {str(e)}")
                self.recording = False
                return False
        return False
        
    def stop_recording(self, engine="whisper"):
        """녹음을 중지하고 남은 텍스트를 반환합니다."""
        print("Stopping recording...")
        try:
            if self.recording:
                self.recording = False
                if engine == "whisper":
                    final_text = self.whisper_transcriber.stop_streaming()
                    if final_text and final_text.strip():
                        self.transcription_buffer.append(final_text.strip())
                        complete_text = " ".join(self.transcription_buffer)
                        print(f"Complete transcription: {complete_text}")
                        return complete_text
                    else:
                        print("No final text received")
            return None
        except Exception as e:
            print(f"Error in stop_recording: {str(e)}")
            return None
        
    def on_transcription_complete(self, text):
        """음성 인식이 완료될 때마다 호출되는 콜백 함수"""
        if text and text.strip():
            print(f"Received transcription: {text.strip()}")
            self.transcription_buffer.append(text.strip())
            # Streamlit의 session_state 업데이트
            if 'consultation_input' in st.session_state:
                current_text = st.session_state.consultation_input
                st.session_state.consultation_input = (current_text + " " + text.strip()).strip()
            else:
                st.session_state.consultation_input = text.strip()
            st.rerun()

class MongoDBManager:
    def __init__(self, uri):
        if not uri:
            raise ValueError("MongoDB URI is required")
        self.client = MongoClient(uri)
        self.db = self.client['divorce_db']
        self.conversations = self.db['conversations']
        self.documents = self.db['document_chunks']
        self.feedback = self.db['feedback']
        self.users = self.db['users']
        
        # 인덱스 생성
        self.documents.create_index([("chunk_hash", 1)], unique=True)
        self.feedback.create_index([("session_id", 1)])
    
    def save_chunk(self, content: str, doc_type: str, embedding: List[float]):
        """문서 청크 저장"""
        chunk_hash = hashlib.md5(content.encode()).hexdigest()
        if not self.documents.find_one({"chunk_hash": chunk_hash}):
            self.documents.insert_one({
                "content": content,
                "doc_type": doc_type,
                "chunk_hash": chunk_hash,
                "embedding": embedding,
                "timestamp": datetime.now()
            })
    
    def save_conversation(self, session_id: str, user_input: str, generated_content: Dict):
        """대화 내용을 MongoDB에 저장"""
        conversation = {
            "session_id": session_id,
            "timestamp": datetime.now(),
            "user_input": user_input,
            "generated_content": generated_content
        }
        self.conversations.insert_one(conversation)
        return conversation

    def get_similar_chunks(self, query_embedding: List[float], doc_type: str, k: int = 3):
        """유사한 청크 검색"""
        raise NotImplementedError("Vector search not implemented in this example.")

    def save_feedback(self, session_id: str, complaint: str, features: List[float], 
                     rating: float, reward: float):
        """강화학습 피드백 저장"""
        try:
            feedback_doc = {
                "session_id": session_id,
                "timestamp": datetime.now(),
                "complaint": complaint,
                "features": features,
                "rating": rating,
                "reward": reward
            }
            result = self.feedback.insert_one(feedback_doc)
            print(f"피드백 저장 성공: {result.inserted_id}")
            return result.inserted_id
        except Exception as e:
            print(f"피드백 저장 실패: {str(e)}")
            raise e

    def get_feedback_statistics(self) -> Dict:
        """피드백 통계 조회"""
        pipeline = [
            {
                "$group": {
                    "_id": None,
                    "avg_rating": {"$avg": "$rating"},
                    "count": {"$sum": 1},
                    "ratings": {"$push": "$rating"}
                }
            }
        ]
        
        stats = list(self.feedback.aggregate(pipeline))
        if stats:
            return {
                "average_rating": stats[0]["avg_rating"],
                "total_feedback": stats[0]["count"],
                "rating_distribution": np.histogram(stats[0]["ratings"], bins=5)[0].tolist()
            }
        return None

class ReinforcementLearner:
    def __init__(self, mongo_db: MongoDBManager):
        self.mongo_db = mongo_db
        self.scaler = StandardScaler()
        self.learning_rate = 0.01
        self.min_feedback_samples = 10  # 최소 피드백 샘플 수
        
    def extract_features(self, complaint: str) -> List[float]:
        """소장에서 특징 추출"""
        features = []
        
        # 1. 문서 길이
        features.append(len(complaint))
        
        # 2. 주요 섹션 포함 여부
        sections = ['청구취지', '청구원인', '입증방법', '첨부서류']
        for section in sections:
            features.append(1.0 if section in complaint else 0.0)
        
        # 3. 법률 용어 사용 빈도
        legal_terms = ['원고', '피고', '위자료', '재산분할', '양육권', '가집행']
        for term in legal_terms:
            features.append(complaint.count(term))
            
        return features
    
    def calculate_reward(self, user_rating: float, complaint_length: int) -> float:
        """보상 계산"""
        # 기본 보상은 사용자 평가
        reward = user_rating
        
        # 문서 길이에 대한 페널티 (너무 짧거나 긴 경우)
        optimal_length = 2000  # 적정 문서 길이
        length_penalty = -abs(complaint_length - optimal_length) / optimal_length
        
        return reward + length_penalty * 0.1  # 길이 페널티의 가중치는 0.1

    def get_best_practices(self) -> Dict[str, any]:
        """높은 평가를 받은 소장들의 특징을 분석"""
        feedback_data = list(self.mongo_db.feedback.find(
            {"rating": {"$gte": 4}}  # 4점 이상 평가받은 소장만 선택
        ).sort("rating", -1).limit(50))  # 상위 50개
        
        if len(feedback_data) < self.min_feedback_samples:
            return None
            
        best_practices = {
            "avg_length": np.mean([len(f["complaint"]) for f in feedback_data]),
            "common_phrases": self._extract_common_phrases(feedback_data),
            "section_patterns": self._analyze_section_patterns(feedback_data),
            "successful_features": self._analyze_successful_features(feedback_data)
        }
        
        return best_practices
        
    def _extract_common_phrases(self, feedback_data: List[Dict]) -> List[str]:
        """성공적인 소장들에서 자주 사용된 문구 추출"""
        # 여기에 문구 추출 로직 구현
        pass
        
    def _analyze_section_patterns(self, feedback_data: List[Dict]) -> Dict:
        """성공적인 소장들의 섹션 구조 분석"""
        # 여기에 섹션 분석 로직 구현
        pass
        
    def _analyze_successful_features(self, feedback_data: List[Dict]) -> Dict:
        """성공적인 소장들의 특징 분석"""
        # 여기에 특징 분석 로직 구현
        pass

class DivorceComplaintGenerator:
    def __init__(self):
        # 환경 변수에서 값 로드
        api_key = os.getenv("OPENAI_API_KEY")
        mongo_uri = os.getenv("MONGO_URI")
        endpoint = os.getenv("LANGCHAIN_ENDPOINT", "https://api.smith.langchain.com")
        langchain_api_key = os.getenv("LANGCHAIN_API_KEY")
        self.project_name = os.getenv("LANGCHAIN_PROJECT", "default_project")  # 기본값 설정

        # LangSmith 설정
        if langchain_api_key:  # API 키가 있을 때만 LangSmith 초기화
            print("LangSmith 클라이언트 초기화")
            self.langsmith_client = Client(
                api_url=endpoint,  # base_url로 변경
                api_key=langchain_api_key,
            )
        else:
            self.langsmith_client = None
            print("LangSmith 클라이언트가 초기화되지 않았습니다. 트레이싱이 비활성화됩니다.")

        # OpenAI Embeddings 초기화
        self.embeddings = OpenAIEmbeddings(api_key=api_key)
        
        # MongoDB 연결
        self.mongo_db = MongoDBManager(mongo_uri)
        self.session_manager = SessionManager()
        self.rl_learner = ReinforcementLearner(self.mongo_db)
        self.user_manager = UserManager(self.mongo_db)
    
    def generate_complaint(self, consultation_text: str) -> dict:
        """소장 생성"""
        run_id = None

        if self.langsmith_client:  # LangSmith가 설정된 경우 실행 생성
            # 실행 생성 (name 및 run_type 포함)
            run_id = self.langsmith_client.create_run(
                name="generate_complaint",  # 실행 이름
                run_type="tool",            # 실행 유형 (예: 'tool', 'chain', 'llm' 등)
                project_name=self.project_name,  # 프로젝트 이름
                inputs={"consultation_text": consultation_text}  # 입력 데이터
            )

        try:
            # 내부 로직 실행
            result = self._generate_complaint_internal(consultation_text)

            # 실행 결과 업데이트
            if self.langsmith_client and run_id:
                self.langsmith_client.update_run(
                    run_id=run_id,
                    outputs=result,
                    status="completed"
                )
            return result
        except Exception as e:
            # 실행 오류 처리
            if self.langsmith_client and run_id:
                self.langsmith_client.update_run(
                    run_id=run_id,
                    error=str(e),
                    status="failed"
                )
            raise e


    def _generate_complaint_internal(self, consultation_text: str) -> dict:
        """실제 소장 생성 로직 (임의로 대체 가능)"""
        claim_chunks = ["Claim data placeholder"]
        relief_chunks = ["Relief data placeholder"]

        return self._generate_with_gpt(consultation_text, claim_chunks, relief_chunks)

    def _generate_with_gpt(self, consultation_text: str, claim_chunks: List[str], relief_chunks: List[str]) -> dict:
        """GPT로 소장 생성 (피드백 학습 적용)"""
        # 성공적인 소장의 특징 가져오기
        best_practices = self.rl_learner.get_best_practices()
        
        if best_practices:
            # 피드백 기반 프롬프트 최적화
            quality_guidelines = f"""
            다음 기준을 충족하는 고품질 소장을 작성해주세요:
            1. 적정 문서 길이: {int(best_practices['avg_length'])} 자 내외
            2. 다음 문구들을 적절히 활용하세요:
               {', '.join(best_practices['common_phrases'][:5])}
            3. 성공적인 소장의 섹션 구조를 따르세요:
               {best_practices['section_patterns']}
            4. 다음 특징들을 반영하세요:
               {best_practices['successful_features']}
            """
        else:
            # 기본 가이드라인 사용
            quality_guidelines = """
            다음 기준을 충족하는 고품질 소장을 작성해주세요:
            1. 모든 필수 섹션을 포함할 것
            2. 구체적이고 명확한 법률 용어 사용
            3. 논리적인 구조와 흐름
            4. 적절한 길이와 상세도
            """
        
        prompt = (
            quality_guidelines + "\n\n" +
            "다음 상담 내용과 참고 문서를 바탕으로 이혼 소장을 작성해주세요.\n\n" +
            f"[상담 내용]\n{consultation_text}\n\n" +
            f"[청구취지 참고문서]\n{' '.join(claim_chunks)}\n\n" +
            f"[청구원인 참고문서]\n{' '.join(relief_chunks)}\n\n"
        )
        
        # GPT 모델에 컨텍스트 추가
        messages = [
            {
                "role": "system",
                "content": "당신은 전문 법률 문서 작성 시스템입니다. 과거의 성공적인 소장 작성 경험을 바탕으로 최적화된 이혼 소장을 작성합니다."
            }
        ]
        
        # 성공적인 예시 추가
        if best_practices and 'example_complaints' in best_practices:
            messages.append({
                "role": "assistant",
                "content": f"다음은 높은 평가를 받은 소장의 예시입니다:\n{best_practices['example_complaints'][0]}"
            })
        
        messages.append({
            "role": "user",
            "content": prompt
        })
        
        response = openai.chat.completions.create(
            model="gpt-4",
            messages=messages,
            temperature=0.3
        )
        
        return response.choices[0].message.content

class DocumentExporter:
    @staticmethod
    def export_to_word(complaint_text: str, filename: str = "이혼소장.docx") -> str:
        """소장을 Word 문서로 변환"""
        doc = Document()
        
        # 문서 스타일 설정
        style = doc.styles['Normal']
        style.font.name = '바탕체'
        style.font.size = Pt(12)
        
        # 제목
        title = doc.add_paragraph()
        title_run = title.add_run("소    장")
        title_run.font.size = Pt(16)
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 본문 추가
        sections = complaint_text.split('\n\n')
        for section in sections:
            p = doc.add_paragraph()
            p.add_run(section)
            
        # 여백 설정
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # 임시 파일로 저장
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            return tmp_file.name

def initialize_session_state():
    """세션 상태 초기화"""
    if 'conversation_history' not in st.session_state:
        st.session_state.conversation_history = []
    if 'generated_complaint' not in st.session_state:
        st.session_state.generated_complaint = None
    if 'edit_mode' not in st.session_state:
        st.session_state.edit_mode = False
    if 'evaluation_mode' not in st.session_state:
        st.session_state.evaluation_mode = False
    if 'rating' not in st.session_state:
        st.session_state.rating = 1  # 최소값 1로 설정
    if 'session_id' not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
    if 'generator' not in st.session_state:
        st.session_state.generator = DivorceComplaintGenerator()
    if 'session_manager' not in st.session_state:
        st.session_state.session_manager = SessionManager()

def display_complaint_actions():
    if st.session_state.generated_complaint:
        st.markdown("---")
        st.subheader("생성된 소장")
        
        # 소장 내용 표시
        if st.session_state.edit_mode:
            # 수정 모드
            edited_complaint = st.text_area(
                "소장 내용 수정",
                value=st.session_state.generated_complaint,
                height=300
            )
            
            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("저장"):
                    st.session_state.generated_complaint = edited_complaint
                    st.session_state.edit_mode = False
                    st.rerun()
            with col2:
                if st.button("취소"):
                    st.session_state.edit_mode = False
                    st.rerun()
        else:
            # 조회 모드
            st.markdown(st.session_state.generated_complaint)
            
            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("수정"):
                    st.session_state.edit_mode = True
                    st.rerun()
            
            # 평가 모드
            if not st.session_state.evaluation_submitted:
                st.markdown("---")
                st.subheader("소장 평가")
                
                # 평점
                rating = st.slider(
                    "소장의 품질은 어떠신가요?",
                    min_value=1,
                    max_value=5,
                    value=st.session_state.get('rating', 3)  # 기본값으로 3 설정
                )
                
                # 피드백
                feedback = st.text_area(
                    "상세 피드백을 남겨주세요",
                    value=st.session_state.get('feedback', ''),
                    height=100
                )
                
                if st.button("평가 제출"):
                    try:
                        current_session = st.session_state.session_manager.get_current_session()
                        if current_session:
                            # 평가 정보 저장
                            st.session_state.rating = rating
                            st.session_state.feedback = feedback
                            
                            # 보상 계산 (예시로 최고 평가를 기준으로 설정)
                            reward = st.session_state.generator.rl_learner.calculate_reward(rating, len(st.session_state.generated_complaint))
                            
                            # MongoDB에 평가 저장
                            st.session_state.generator.mongo_db.save_feedback(
                                current_session.session_id,
                                st.session_state.generated_complaint,
                                st.session_state.generator.rl_learner.extract_features(st.session_state.generated_complaint),
                                rating,
                                reward
                            )
                            
                            # 세션 상태 업데이트
                            st.session_state.evaluation_submitted = True
                            st.success("평가가 저장되었습니다. 감사합니다!")
                            st.rerun()
                        else:
                            st.error("현재 세션을 찾을 수 없습니다.")
                    except Exception as e:
                        st.error(f"평가 저장 중 오류가 발생했습니다: {str(e)}")
            else:
                st.info("이미 평가를 제출하셨습니다. 감사합니다!")

def main():
    st.set_page_config(
        page_title="법률 상담 도우미",
        page_icon="⚖️",
        layout="wide"
    )
    
    # 초기화
    if 'session_id' not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())

    if "stt_manager" not in st.session_state:
        st.session_state.stt_manager = STTManager()

    if "generator" not in st.session_state:
        st.session_state.generator = DivorceComplaintGenerator()
        
    if "session_manager" not in st.session_state:
        st.session_state.session_manager = SessionManager()

    if "generated_complaint" not in st.session_state:
        st.session_state.generated_complaint = None

    if "recording" not in st.session_state:
        st.session_state.recording = False
        
    if "consultation_text" not in st.session_state:
        st.session_state.consultation_text = ""
        
    if "edit_mode" not in st.session_state:
        st.session_state.edit_mode = False
        
    if "evaluation_mode" not in st.session_state:
        st.session_state.evaluation_mode = False
        
    if "conversation_history" not in st.session_state:
        st.session_state.conversation_history = []
        
    # 평가 관련 상태 초기화
    if "rating" not in st.session_state:
        st.session_state.rating = 1  # 최소값 1로 설정
        
    if "feedback" not in st.session_state:
        st.session_state.feedback = ""
        
    if "evaluation_submitted" not in st.session_state:
        st.session_state.evaluation_submitted = False
    
    # 2분할 레이아웃
    left_col, right_col = st.columns([2, 8])
    
    # 왼쪽 사이드바 - 세션 목록
    with left_col:
        if st.session_state.generator.user_manager.is_logged_in():
            display_sessions()
    
    # 오른쪽 컬럼 - 메인 컨텐츠
    with right_col:
        st.title("히어로 법률 도우미")
        
        # 로그인/회원가입 섹션
        if not st.session_state.generator.user_manager.is_logged_in():
            tab1, tab2 = st.tabs(["로그인", "회원가입"])
            
            with tab1:
                username = st.text_input("사용자 이름", key="login_username")
                password = st.text_input("비밀번호", type="password", key="login_password")
                if st.button("로그인"):
                    if st.session_state.generator.user_manager.login(username, password):
                        st.success("로그인 성공")
                        st.rerun()
                    else:
                        st.error("로그인 실패")
            
            with tab2:
                reg_username = st.text_input("사용자 이름", key="reg_username")
                reg_password = st.text_input("비밀번호", type="password", key="reg_password")
                reg_password_confirm = st.text_input("비밀번호 확인", type="password")
                
                if st.button("회원가입"):
                    if reg_password != reg_password_confirm:
                        st.error("비밀번호가 일치하지 않습니다.")
                    else:
                        success, message = st.session_state.generator.user_manager.register(reg_username, reg_password)
                        if success:
                            st.success(message)
                            # 자동 로그인
                            if st.session_state.generator.user_manager.login(reg_username, reg_password):
                                st.rerun()
                        else:
                            st.error(message)
        
        # 로그인된 경우에만 메인 기능 표시
        if st.session_state.generator.user_manager.is_logged_in():
            # 상담 입력
            consultation_text = st.text_area(
                "상담 내용을 입력하세요",
                value=st.session_state.get("consultation_text", ""),
                placeholder="상담 내용을 입력해주세요...",
                height=200
            )
            
            # 음성 입력 컨트롤
            col1, col2 = st.columns([1, 1])
            with col1:
                stt_engine = st.selectbox(
                    "음성 인식 엔진",
                    ["Whisper", "GCP"],
                    key="stt_engine",
                    label_visibility="collapsed"  # 레이블 숨기기
                )
            with col2:
                # 녹음 상태 표시
                if st.session_state.recording:
                    st.info("🎙️ 녹음 중...")
                
                # 녹음 버튼
                button_text = "⏹️ 녹음 중지" if st.session_state.recording else "🎤 녹음 시작"
                if st.button(button_text, use_container_width=True):  # 버튼 너비를 컨테이너에 맞춤
                    if not st.session_state.recording:
                        if st.session_state.stt_manager.start_recording(engine=stt_engine.lower()):
                            st.session_state.recording = True
                            print("RECORDING START")
                            st.rerun()
                    else:
                        print("RECORDING STOP")
                        st.info("변환 중...")
                        transcribed_text = st.session_state.stt_manager.stop_recording()
                        print("Transcribing to STT...")
                        if transcribed_text:
                            print(f"Transcribed text: {transcribed_text}")
                            if not st.session_state.consultation_text:
                                st.session_state.consultation_text = transcribed_text
                            else:
                                st.session_state.consultation_text += " " + transcribed_text
                        st.session_state.recording = False
                        st.rerun()
            
            # 소장 생성 버튼
            if st.button("소장 생성", type="primary"):
                if consultation_text:
                    with st.spinner("소장을 생성하고 있습니다..."):
                        try:
                            # 소장 생성
                            complaint = st.session_state.generator.generate_complaint(consultation_text)
                            st.session_state.generated_complaint = complaint
                            
                            # 세션 기록과 MongoDB에 저장
                            st.session_state.session_manager.save_session(consultation_text, complaint)
                            st.session_state.generator.mongo_db.save_conversation(
                                st.session_state.session_manager.get_sessions()[-1].session_id,
                                consultation_text,
                                complaint
                            )
                            
                            st.rerun()
                        except Exception as e:
                            st.error(f"오류가 발생했습니다: {str(e)}")
                else:
                    st.warning("상담 내용을 입력해주세요.")

            # 소장 평가 및 수정 UI 표시
            display_complaint_actions()

if __name__ == "__main__":
    main()