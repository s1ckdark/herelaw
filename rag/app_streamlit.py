import streamlit as st
import openai
import os
from dotenv import load_dotenv
from langchain_community.vectorstores import FAISS
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.llms import OpenAI
from pymongo import MongoClient
from datetime import datetime
import uuid
import hashlib
from sklearn.preprocessing import StandardScaler
import numpy as np
from docx import Document
import time
import base64
import tempfile
import os

from langsmith import Client
from typing import List, Dict, Optional
import json
import sys
sys.path.append("../stt")
from gcpApi import record_audio as gcp_record
from gcpApi import audio_to_text as gcp_transcribe
from whisperApi import AudioTranscriber
import extra_streamlit_components as stx
from datetime import timedelta
import requests

# .env 파일 로드
load_dotenv()

def get_database():
    """MongoDB 데이터베이스 연결을 반환합니다."""
    # MongoDB 연결 정보
    MONGODB_URI = os.getenv('MONGO_URI')  # .env 파일의 MONGO_URI 사용
    MONGODB_DB = os.getenv('MONGODB_DB', 'herelaw')
    
    # MongoDB 클라이언트 생성
    client = MongoClient(MONGODB_URI)
    return client[MONGODB_DB]

class Session:
    def __init__(self, session_id, timestamp, consultation_text="", generated_content="", rating=None, feedback=""):
        self.session_id = session_id
        self.timestamp = timestamp
        self.consultation_text = consultation_text
        self.generated_content = generated_content
        self.rating = rating
        self.feedback = feedback

class UserManager:
    def __init__(self, mongo_db):
        self.mongo_db = mongo_db
        if 'user' not in st.session_state:
            st.session_state.user = None
        
        # 쿠키 매니저 초기화 (고유한 키 사용)
        self.cookie_manager = stx.CookieManager(key="user_manager")
        
        # 쿠키에서 세션 복원
        self._restore_session_from_cookie()

    def _restore_session_from_cookie(self):
        """쿠키에서 사용자 세션을 복원합니다."""
        try:
            if not st.session_state.user:  # 세션에 사용자가 없을 때만 쿠키 확인
                username = self.cookie_manager.get("username")
                if username:
                    user = self.mongo_db.db.users.find_one({"username": username})
                    if user:
                        st.session_state.user = {
                            "username": user["username"],
                            "user_id": str(user["_id"])
                        }
        except Exception as e:
            print(f"Error restoring session from cookie: {str(e)}")

    def login(self, username, password):
        try:
            user = self.mongo_db.db.users.find_one({"username": username, "password": password})
            if user:
                st.session_state.user = {
                    "username": user["username"],
                    "user_id": str(user["_id"])
                }
                # 쿠키 설정 (30일 유효기간)
                try:
                    self.cookie_manager.set("username", username, expires_at=datetime.now() + timedelta(days=30))
                except Exception as e:
                    print(f"Error setting cookie: {str(e)}")
                return True
            return False
        except Exception as e:
            print(f"Error in login: {str(e)}")
            return False

    def logout(self):
        try:
            st.session_state.user = None
            st.session_state.sessions = []
            st.session_state.current_session_id = None
            # 쿠키 삭제
            try:
                self.cookie_manager.delete("username")
            except Exception as e:
                print(f"Error deleting cookie: {str(e)}")
        except Exception as e:
            print(f"Error in logout: {str(e)}")

    def is_logged_in(self):
        if not st.session_state.user:
            self._restore_session_from_cookie()
        return st.session_state.user is not None

    def get_current_user(self):
        if not st.session_state.user:
            self._restore_session_from_cookie()
        return st.session_state.user

    def register(self, username, password):
        try:
            # 이미 존재하는 사용자인지 확인
            existing_user = self.mongo_db.db.users.find_one({"username": username})
            if existing_user:
                return False, "이미 존재하는 사용자입니다."
            
            # 새 사용자 등록
            user = {
                "username": username,
                "password": password,  # 실제 구현시 해시 처리 필요
                "created_at": datetime.now()
            }
            self.mongo_db.db.users.insert_one(user)
            return True, "사용자 등록이 완료되었습니다."
        except Exception as e:
            return False, f"사용자 등록 중 오류가 발생했습니다: {str(e)}"

class SessionManager:
    def __init__(self):
        if 'sessions' not in st.session_state:
            st.session_state.sessions = []
        if 'current_session_id' not in st.session_state:
            st.session_state.current_session_id = None
        self.db = get_database()
        self.sessions_collection = self.db['sessions']

    def save_session(self, consultation_text, generated_content):
        user = st.session_state.get('user')
        if not user:
            raise ValueError("User not logged in")

        session = Session(
            session_id=str(uuid.uuid4()),
            timestamp=datetime.now(),
            consultation_text=consultation_text,
            generated_content=generated_content
        )
        session.user_id = user['user_id']
        session.username = user['username']
        
        # MongoDB에 저장
        session_data = {
            'session_id': session.session_id,
            'user_id': session.user_id,
            'username': session.username,
            'timestamp': session.timestamp,
            'consultation_text': session.consultation_text,
            'generated_content': session.generated_content,
            'rating': None,
            'feedback': ''
        }
        self.sessions_collection.insert_one(session_data)
        
        st.session_state.sessions.append(session)
        st.session_state.current_session_id = session.session_id
        return session.session_id

    def get_sessions(self):
        """현재 로그인한 사용자의 세션을 MongoDB에서 가져옵니다."""
        user = st.session_state.get('user')
        if not user:
            return []
        
        # MongoDB에서 세션 로드
        session_docs = self.sessions_collection.find(
            {'user_id': user['user_id']},
            sort=[('timestamp', -1)]  # 최신 순으로 정렬
        )
        
        sessions = []
        for doc in session_docs:
            session = Session(
                session_id=doc['session_id'],
                timestamp=doc['timestamp'],
                consultation_text=doc['consultation_text'],
                generated_content=doc['generated_content']
            )
            session.user_id = doc['user_id']
            session.username = doc['username']
            session.rating = doc.get('rating')
            session.feedback = doc.get('feedback', '')
            sessions.append(session)
        
        # 세션 상태 업데이트
        st.session_state.sessions = sessions
        return sessions

    def update_current_session(self, rating=None, feedback=None):
        if not st.session_state.current_session_id:
            return False
        
        update_data = {}
        if rating is not None:
            update_data['rating'] = rating
        if feedback is not None:
            update_data['feedback'] = feedback
        
        if update_data:
            # MongoDB 업데이트
            self.sessions_collection.update_one(
                {'session_id': st.session_state.current_session_id},
                {'$set': update_data}
            )
            
            # 메모리 상의 세션도 업데이트
            for session in st.session_state.sessions:
                if session.session_id == st.session_state.current_session_id:
                    if rating is not None:
                        session.rating = rating
                    if feedback is not None:
                        session.feedback = feedback
                    return True
        return False

class STTManager:
    def __init__(self, openai_api_key=None):
        """
        STT(Speech-to-Text) 관리자 클래스 초기화
        
        Args:
            openai_api_key (str, optional): OpenAI API 키
        """
        self.openai_api_key = openai_api_key

    def convert_webm_to_wav(self, webm_path, wav_path):
        """
        WebM 파일을 WAV 형식으로 변환합니다.
        
        Args:
            webm_path (str): WebM 파일 경로
            wav_path (str): 저장할 WAV 파일 경로
        """
        from moviepy.editor import AudioFileClip
        
        audio = AudioFileClip(webm_path)
        audio.write_audiofile(wav_path)
        audio.close()

    def transcribe_with_whisper(self, audio_path):
        """
        오디오 파일을 Whisper API를 사용하여 텍스트로 변환합니다.
        
        Args:
            audio_path (str): 오디오 파일 경로
        
        Returns:
            str: 변환된 텍스트
        """
        import openai
        
        if not self.openai_api_key:
            raise ValueError("OpenAI API key is required for Whisper transcription")
        
        openai.api_key = self.openai_api_key
        
        with open(audio_path, "rb") as audio_file:
            transcript = openai.Audio.transcribe(
                model="whisper-1",
                file=audio_file
            )
        
        return transcript["text"]

    def process_webm_to_text(self, webm_path):
        """
        WebM 파일을 WAV로 변환하고 텍스트로 변환합니다.
        
        Args:
            webm_path (str): WebM 파일 경로
        
        Returns:
            str: 변환된 텍스트
        """
        import os
        
        # WAV 파일 경로 생성
        wav_path = os.path.splitext(webm_path)[0] + ".wav"
        
        try:
            # WebM을 WAV로 변환
            self.convert_webm_to_wav(webm_path, wav_path)
            
            # WAV 파일을 텍스트로 변환
            text = self.transcribe_with_whisper(wav_path)
            
            return text
            
        except Exception as e:
            raise e
            
        finally:
            # 임시 WAV 파일 삭제
            if os.path.exists(wav_path):
                os.remove(wav_path)

class MongoDBManager:
    def __init__(self, uri):
        if not uri:
            raise ValueError("MongoDB URI is required")
        self.client = MongoClient(uri)
        self.db = self.client['divorce_db']
        self.conversations = self.db['conversations']
        self.documents = self.db['document_chunks']
        self.feedback = self.db['feedback']
        self.users = self.db['users']
        
        # 인덱스 생성
        self.documents.create_index([("chunk_hash", 1)], unique=True)
        self.feedback.create_index([("session_id", 1)])
    
    def save_chunk(self, content: str, doc_type: str, embedding: List[float]):
        chunk_hash = hashlib.md5(content.encode()).hexdigest()
        if not self.documents.find_one({"chunk_hash": chunk_hash}):
            self.documents.insert_one({
                "content": content,
                "doc_type": doc_type,
                "chunk_hash": chunk_hash,
                "embedding": embedding,
                "timestamp": datetime.now()
            })
    
    def save_conversation(self, session_id: str, user_input: str, generated_content: Dict):
        conversation = {
            "session_id": session_id,
            "timestamp": datetime.now(),
            "user_input": user_input,
            "generated_content": generated_content
        }
        self.conversations.insert_one(conversation)
        return conversation

    def get_similar_chunks(self, query_embedding: List[float], doc_type: str, k: int = 3):
        raise NotImplementedError("Vector search not implemented in this example.")

    def save_feedback(self, session_id: str, complaint: str, features: List[float], 
                     rating: float, reward: float):
        """강화학습 피드백 저장"""
        try:
            feedback_doc = {
                "session_id": session_id,
                "timestamp": datetime.now(),
                "complaint": complaint,
                "features": features,
                "rating": rating,
                "reward": reward
            }
            result = self.feedback.insert_one(feedback_doc)
            print(f"피드백 저장 성공: {result.inserted_id}")
            return result.inserted_id
        except Exception as e:
            print(f"피드백 저장 실패: {str(e)}")
            raise e

    def get_feedback_statistics(self) -> Dict:
        """피드백 통계 조회"""
        pipeline = [
            {
                "$group": {
                    "_id": None,
                    "avg_rating": {"$avg": "$rating"},
                    "count": {"$sum": 1},
                    "ratings": {"$push": "$rating"}
                }
            }
        ]
        
        stats = list(self.feedback.aggregate(pipeline))
        if stats:
            return {
                "average_rating": stats[0]["avg_rating"],
                "total_feedback": stats[0]["count"],
                "rating_distribution": np.histogram(stats[0]["ratings"], bins=5)[0].tolist()
            }
        return None

class ReinforcementLearner:
    def __init__(self, mongo_db: MongoDBManager):
        self.mongo_db = mongo_db
        self.scaler = StandardScaler()
        self.learning_rate = 0.01
        self.min_feedback_samples = 10  # 최소 피드백 샘플 수
        
    def extract_features(self, complaint: str) -> List[float]:
        """소장에서 특징 추출"""
        features = []
        
        # 1. 문서 길���
        features.append(len(complaint))
        
        # 2. 주요 섹션 포함 여부
        sections = ['청구취지', '청구원인', '입증방법', '첨부서류']
        for section in sections:
            features.append(1.0 if section in complaint else 0.0)
        
        # 3. 법률 용어 사용 빈도
        legal_terms = ['원고', '피고', '위자료', '재산분할', '양육권', '가집행']
        for term in legal_terms:
            features.append(complaint.count(term))
            
        return features
    
    def calculate_reward(self, user_rating: float, complaint_length: int) -> float:
        """보상 계산"""
        # 기본 보상은 사용자 평가
        reward = user_rating
        
        # 문서 길이에 대한 페널티 (너무 짧거나 긴 경우)
        optimal_length = 2000  # 적정 문서 길이
        length_penalty = -abs(complaint_length - optimal_length) / optimal_length
        
        return reward + length_penalty * 0.1  # 길이 페널티의 가중치는 0.1

    def get_best_practices(self) -> Dict[str, any]:
        """높은 평가를 받은 소장들의 특징을 분석"""
        feedback_data = list(self.mongo_db.feedback.find(
            {"rating": {"$gte": 4}}  # 4점 이상 평가받은 소장만 선택
        ).sort("rating", -1).limit(50))  # 상위 50개
        
        if len(feedback_data) < self.min_feedback_samples:
            return None
            
        best_practices = {
            "avg_length": np.mean([len(f["complaint"]) for f in feedback_data]),
            "common_phrases": self._extract_common_phrases(feedback_data),
            "section_patterns": self._analyze_section_patterns(feedback_data),
            "successful_features": self._analyze_successful_features(feedback_data)
        }
        
        return best_practices
        
    def _extract_common_phrases(self, feedback_data: List[Dict]) -> List[str]:
        """성공적인 소장들에서 자주 사용된 문구 추출"""
        # 여기에 문구 추출 로직 구현
        pass
        
    def _analyze_section_patterns(self, feedback_data: List[Dict]) -> Dict:
        """성공적인 소장들의 섹션 구조 분석"""
        # 여기에 섹션 분석 로직 구현
        pass
        
    def _analyze_successful_features(self, feedback_data: List[Dict]) -> Dict:
        """성공적인 소장들의 특징 분석"""
        # 여기에 특징 분석 로직 구현
        pass

class DivorceComplaintGenerator:
    def __init__(self):
        # 환 변수에서 값 로드
        api_key = os.getenv("OPENAI_API_KEY")
        mongo_uri = os.getenv("MONGO_URI")
        endpoint = os.getenv("LANGCHAIN_ENDPOINT", "https://api.smith.langchain.com")
        langchain_api_key = os.getenv("LANGCHAIN_API_KEY")
        self.project_name = os.getenv("LANGCHAIN_PROJECT", "default_project")  # 기본값 설정

        # LangSmith 설정
        if langchain_api_key:  # API 키가 있을 때만 LangSmith 초기화
            print("LangSmith 클라이언트 초기화")
            self.langsmith_client = Client(
                api_url=endpoint,  # base_url로 변경
                api_key=langchain_api_key,
            )
        else:
            self.langsmith_client = None
            print("LangSmith 클라이언트가 초기화되지 않았습니다. 트레이싱이 비활성화됩니다.")

        # OpenAI Embeddings 초기화
        self.embeddings = OpenAIEmbeddings(api_key=api_key)
        
        # MongoDB 연결
        self.mongo_db = MongoDBManager(mongo_uri)
        self.session_manager = SessionManager()
        self.rl_learner = ReinforcementLearner(self.mongo_db)
        self.user_manager = UserManager(self.mongo_db)
        self.stt_manager = STTManager(api_key)
    
    def generate_complaint(self, consultation_text: str) -> dict:
        """소장 생성"""
        run_id = None

        if self.langsmith_client:  # LangSmith가 설정된 경우 실행 생성
            # 실행 생성 (name 및 run_type 포함)
            run_id = self.langsmith_client.create_run(
                name="generate_complaint",  # 실행 이름
                run_type="tool",            # 실행 유형 (예: 'tool', 'chain', 'llm' 등)
                project_name=self.project_name,  # 프로젝트 이름
                inputs={"consultation_text": consultation_text}  # 입력 데이터
            )

        try:
            # 내부 로직 실행
            result = self._generate_complaint_internal(consultation_text)

            # 실행 결과 업데이트
            if self.langsmith_client and run_id:
                self.langsmith_client.update_run(
                    run_id=run_id,
                    outputs=result,
                    status="completed"
                )
            return result
        except Exception as e:
            # 실행 오류 처리
            if self.langsmith_client and run_id:
                self.langsmith_client.update_run(
                    run_id=run_id,
                    error=str(e),
                    status="failed"
                )
            raise e


    def _generate_complaint_internal(self, consultation_text: str) -> dict:
        """실제 소장 생성 로직 (임의로 대체 가능)"""
        claim_chunks = ["Claim data placeholder"]
        relief_chunks = ["Relief data placeholder"]

        return self._generate_with_gpt(consultation_text, claim_chunks, relief_chunks)

    def _generate_with_gpt(self, consultation_text: str, claim_chunks: List[str], relief_chunks: List[str]) -> dict:
        """GPT로 소장 생성 (피드백 학습 적용)"""
        # 성공적인 소장의 특징 가져오기
        best_practices = self.rl_learner.get_best_practices()
        
        if best_practices:
            # 피드백 기반 프롬프트 최적화
            quality_guidelines = f"""
            다음 기준을 충족하는 고품질 소장을 작성해주세요:
            1. 적정 문서 길이: {int(best_practices['avg_length'])} 자 내외
            2. 다음 문구들을 적절히 활용하세요:
               {', '.join(best_practices['common_phrases'][:5])}
            3. 성공적인 소장의 섹션 구조를 따르세요:
               {best_practices['section_patterns']}
            4. 다음 특징들을 반영하세요:
               {best_practices['successful_features']}
            """
        else:
            # 기본 가이드라인 사용
            quality_guidelines = """
            다음 기준을 충족하는 고품질 소장을 작성해주세요:
            1. 모든 필수 섹션을 포함할 것
            2. 구체적이고 명확한 법률 용어 사용
            3. 논리적인 구조와 흐름
            4. 적절한 길이와 상세도
            """
        
        prompt = (
            quality_guidelines + "\n\n" +
            "다음 상담 내용과 참고 문서를 바탕으로 이혼 소장을 작성해주세요.\n\n" +
            f"[상담 내용]\n{consultation_text}\n\n" +
            f"[청구취지 참고문서]\n{' '.join(claim_chunks)}\n\n" +
            f"[청구원인 참고문서]\n{' '.join(relief_chunks)}\n\n"
        )
        
        # GPT 모델에 컨텍스트 추가
        messages = [
            {
                "role": "system",
                "content": "당신은 전문 법률 문서 작성 시스템입니다. 과거의 성공적인 소장 작성 경험을 바탕으로 최적화된 이혼 소장을 작성합니다."
            }
        ]
        
        # 성공적인 예시 추가
        if best_practices and 'example_complaints' in best_practices:
            messages.append({
                "role": "assistant",
                "content": f"다음은 높은 평가를 받은 소장의 예시입니다:\n{best_practices['example_complaints'][0]}"
            })
        
        messages.append({
            "role": "user",
            "content": prompt
        })
        
        response = openai.chat.completions.create(
            model="gpt-4",
            messages=messages,
            temperature=0.3
        )
        
        return response.choices[0].message.content

class DocumentExporter:
    @staticmethod
    def export_to_word(complaint_text: str, filename: str = "이혼소장.docx") -> str:
        """소장을 Word 문서로 변환"""
        doc = Document()
        
        # 문서 스타일 설정
        style = doc.styles['Normal']
        style.font.name = '바탕체'
        style.font.size = Pt(12)
        
        # 제목
        title = doc.add_paragraph()
        title_run = title.add_run("소    장")
        title_run.font.size = Pt(16)
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 본문 추
        sections = complaint_text.split('\n\n')
        for section in sections:
            p = doc.add_paragraph()
            p.add_run(section)
            
        # 여백 설정
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # 임시 파일로 저장
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            return tmp_file.name

def initialize_session_state():
    """세션 상태 초기화"""
    if 'conversation_history' not in st.session_state:
        st.session_state.conversation_history = []
    if 'generated_complaint' not in st.session_state:
        st.session_state.generated_complaint = None
    if 'edit_mode' not in st.session_state:
        st.session_state.edit_mode = False
    if 'evaluation_mode' not in st.session_state:
        st.session_state.evaluation_mode = False
    if 'rating' not in st.session_state:
        st.session_state.rating = 1  # 최소값 1로 설정
    if 'session_id' not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
    if 'generator' not in st.session_state:
        st.session_state.generator = DivorceComplaintGenerator()
    if 'session_manager' not in st.session_state:
        st.session_state.session_manager = SessionManager()

def display_sessions():
    """Display the list of sessions recorded by the user.

    This function displays the list of sessions recorded by the user.
    It also provides a button to start a new session and a button to use a specific session.

    Parameters
    ----------
    None

    Returns
    -------
    None
    """
    if not st.session_state.user:
        return

    st.write("### 💬 상담 기록")
    
    # Display the current user information and the logout button
    col1, col2 = st.columns([2, 1])
    with col1:
        st.markdown(f"**👤 현재 사용자:** {st.session_state.user['username']}")
    with col2:
        if st.button("로그아웃", use_container_width=True):
            st.session_state.generator.user_manager.logout()
            st.rerun()
    
    st.markdown("---")  # Add a separator
    
    if st.button("✨ 새로운 상담 시작하기", type="primary", use_container_width=True):
        st.session_state.current_session_id = None
        st.session_state.consultation_text = ""
        st.session_state.generated_complaint = ""
        st.rerun()

    # Get the list of sessions
    sessions = st.session_state.session_manager.get_sessions()
    
    if not sessions:
        st.info("아직 상담 기록이 없습니다.")
        return

    # Display the list of sessions
    for idx, session in enumerate(reversed(sessions)):
        with st.expander(f"상담 {idx + 1} - {session.timestamp.strftime('%Y-%m-%d %H:%M')}"):
            st.markdown("#### 상담 내용")
            st.write(session.consultation_text)
            
            st.markdown("#### 생성된 소장")
            st.write(session.generated_content)
            
            # Display the evaluation section
            st.markdown("#### 평가")
            if hasattr(session, 'rating') and session.rating:
                st.write(f"⭐ 평점: {session.rating}")
            if hasattr(session, 'feedback') and session.feedback:
                st.write(f"💭 피드백: {session.feedback}")
            
            # Display the "Use this version" button if the current session is not the same as the session being displayed
            if session.session_id != st.session_state.current_session_id:
                if st.button("이 버전 사용", key=f"use_version_{session.session_id}", use_container_width=True):
                    st.session_state.current_session_id = session.session_id
                    st.rerun()

def display_complaint_actions():
    if st.session_state.generated_complaint:
        st.markdown("---")
        st.subheader("생성된 소장")
        
        # 소장 내용 표시
        if st.session_state.edit_mode:
            # 수정 모드
            edited_complaint = st.text_area(
                "소장 내용 수정",
                value=st.session_state.generated_complaint,
                height=300
            )
            
            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("저장"):
                    st.session_state.generated_complaint = edited_complaint
                    st.session_state.edit_mode = False
                    st.rerun()
            with col2:
                if st.button("취소"):
                    st.session_state.edit_mode = False
                    st.rerun()
        else:
            # 조회 모드
            st.markdown(st.session_state.generated_complaint)
            
            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("수정"):
                    st.session_state.edit_mode = True
                    st.rerun()
            
            # 평가 모드
            if not st.session_state.evaluation_submitted:
                st.markdown("---")
                st.subheader("소장 평가")
                
                # 평점
                rating = st.slider(
                    "소장의 품질은 어떠신가요?",
                    min_value=1,
                    max_value=5,
                    value=st.session_state.get('rating', 3)  # 기본값으로 3 설정
                )
                
                # 피드백
                feedback = st.text_area(
                    "상세 피드백을 남겨주세요",
                    value=st.session_state.get('feedback', ''),
                    height=100
                )
                
                if st.button("평가 제출"):
                    try:
                        current_session = st.session_state.session_manager.get_current_session()
                        if current_session:
                            # 평가 정보 저장
                            st.session_state.rating = rating
                            st.session_state.feedback = feedback
                            
                            # 보상 계산 (예시로 최고 평가를 기준으로 설정)
                            reward = st.session_state.generator.rl_learner.calculate_reward(rating, len(st.session_state.generated_complaint))
                            
                            # MongoDB에 평가 저장
                            st.session_state.generator.mongo_db.save_feedback(
                                current_session.session_id,
                                st.session_state.generated_complaint,
                                st.session_state.generator.rl_learner.extract_features(st.session_state.generated_complaint),
                                rating,
                                reward
                            )
                            
                            # 세션 상태 업데이트
                            st.session_state.evaluation_submitted = True
                            st.success("평가가 저장되었습니다. 감사합니다!")
                            st.rerun()
                        else:
                            st.error("현재 세션을 찾을 수 없습니다.")
                    except Exception as e:
                        st.error(f"평가 저장 중 오류가 발생했습니다: {str(e)}")
            else:
                st.info("이미 평가를 제출하셨습니다. 감사합니다!")

def main():
    st.set_page_config(
        page_title="법률 상담 도우미",
        page_icon="⚖️",
        layout="wide"
    )
    
    # 초기화
    if 'session_id' not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())

    if "stt_manager" not in st.session_state:
        st.session_state.stt_manager = STTManager()

    if "generator" not in st.session_state:
        st.session_state.generator = DivorceComplaintGenerator()
        
    if "session_manager" not in st.session_state:
        st.session_state.session_manager = SessionManager()

    if "generated_complaint" not in st.session_state:
        st.session_state.generated_complaint = None

    if "recording" not in st.session_state:
        st.session_state.recording = False
        
    if "consultation_text" not in st.session_state:
        st.session_state.consultation_text = ""
        
    if "edit_mode" not in st.session_state:
        st.session_state.edit_mode = False
        
    if "evaluation_mode" not in st.session_state:
        st.session_state.evaluation_mode = False
        
    if "conversation_history" not in st.session_state:
        st.session_state.conversation_history = []
        
    # 평가 관련 상태 초기화
    if "rating" not in st.session_state:
        st.session_state.rating = 1  # 최소값 1로 설정
        
    if "feedback" not in st.session_state:
        st.session_state.feedback = ""
        
    if "evaluation_submitted" not in st.session_state:
        st.session_state.evaluation_submitted = False
    
    # 오디오 처리 라우트
    if 'process_audio' in st.query_params:
        audio_file = st.files['audio']
        timestamp = st.form_data['timestamp']
        
        if audio_file is not None:
            stt_manager = STTManager()
            result = stt_manager.process_webm_to_text(audio_file.name)
            
            if result:
                st.json({
                    'status': 'success',
                    'text': result
                })
            else:
                st.json({
                    'status': 'error',
                    'message': '오디오 처리 실패'
                })
        return
    
    # 2분할 레이아웃
    left_col, right_col = st.columns([2, 8])
    
    # 왼쪽 사이드바 - 세션 목록
    with left_col:
        if st.session_state.generator.user_manager.is_logged_in():
            display_sessions()
    
    # 오른쪽 컬럼 - 메인 컨텐츠
    with right_col:
        st.title("히어로 법률 도우미")
               
        # 로그인/회원가입 섹션
        if not st.session_state.generator.user_manager.is_logged_in():
            # 쿠키에서 username 확인
            auth_cookie_manager = stx.CookieManager(key="auth_cookie")
            username_cookie = auth_cookie_manager.get("username")
            
            if username_cookie:
                # username 쿠키가 있으면 자동 로그인 시도
                user = st.session_state.generator.user_manager.mongo_db.db.users.find_one({"username": username_cookie})
                if user:
                    st.session_state.generator.user_manager.login(username_cookie, user["password"])
                    st.rerun()
            else:
                # username 쿠키가 없을 때만 로그인 폼 표시
                tab1, tab2 = st.tabs(["로그인", "회원가입"])
                
                with tab1:
                    username = st.text_input("사용자 이름", key="login_username")
                    password = st.text_input("비밀번호", type="password", key="login_password")
                    
                    if st.button("로그인"):
                        if st.session_state.generator.user_manager.login(username, password):
                            # 로그인 성공 시 username을 쿠키에 저장
                            auth_cookie_manager.set("username", username, expires_at=datetime.now() + timedelta(days=30))
                            st.success("로그인되었습니다!")
                            st.rerun()
                        else:
                            st.error("로그인에 실패했습니다. 사용자 이름과 비밀번��를 확인해주세요.")
                
                with tab2:
                    new_username = st.text_input("사용자 이름", key="register_username")
                    new_password = st.text_input("비밀번호", type="password", key="register_password")
                    new_password_confirm = st.text_input("비밀번호 확인", type="password", key="register_password_confirm")
                    
                    if st.button("회원가입"):
                        if new_password != new_password_confirm:
                            st.error("비밀번호가 일치하지 않습니다.")
                        else:
                            success, message = st.session_state.generator.user_manager.register(new_username, new_password)
                            if success:
                                st.success(message)
                                # 회원가입 성공 시 자동 로그인
                                if st.session_state.generator.user_manager.login(new_username, new_password):
                                    # username을 쿠키에 저장
                                    auth_cookie_manager.set("username", new_username, expires_at=datetime.now() + timedelta(days=30))
                                    st.rerun()
                            else:
                                st.error(message)
        
        # 로그인된 경우에만 메인 기능 표시
        if st.session_state.generator.user_manager.is_logged_in():
            # 상담 입력
            consultation_text = st.text_area(
                "상담 내용을 입력하세요",
                value=st.session_state.get("consultation_text", ""),
                placeholder="상담 내용을 입력해주세요...",
                height=200
            )
            
            # 음성 입력 컨트롤
            col1, col2 = st.columns([1, 1])
            with col1:
                stt_engine = st.selectbox(
                    "음성 인식 엔진 선택",
                    ["Whisper", "GCP"],
                    key="stt_engine"
                )
            
            # 마이크 입력 처리
            recording_html = """
                <div style="padding: 10px; border: 1px solid #ddd; border-radius: 5px;">
                    <button id="recordToggle" onclick="toggleRecording()"
                        style="background-color: #4CAF50; color: white; padding: 10px 20px; border: none; border-radius: 4px;">
                        녹음 시작
                    </button>
                    <audio id="audioPlayback" controls style="display: none; margin-top: 10px; width: 100%;"></audio>
                    <div id="recordingStatus" style="margin-top: 10px; color: #666;"></div>
                    <div id="errorMessage" style="margin-top: 10px; color: #f44336;"></div>
                    <div id="audioLog" style="margin-top: 10px; color: #666; font-size: 0.8em;"></div>
                </div>
                
                <script>
                    let mediaRecorder;
                    let audioChunks = [];
                    let isRecording = false;
                    let audioLogInterval;
                    
                    const statusDiv = document.getElementById('recordingStatus');
                    const errorDiv = document.getElementById('errorMessage');
                    const audioLog = document.getElementById('audioLog');
                    const recordButton = document.getElementById('recordToggle');
                    
                    async function toggleRecording() {
                        try {
                            if (!isRecording) {
                                const stream = await navigator.mediaDevices.getUserMedia({ audio: true });
                                mediaRecorder = new MediaRecorder(stream, {
                                    mimeType: 'audio/webm;codecs=opus'
                                });
                                
                                audioChunks = [];
                                
                                mediaRecorder.ondataavailable = (event) => {
                                    audioChunks.push(event.data);
                                };
                                
                                mediaRecorder.onstop = async () => {
                                    try {
                                        clearInterval(audioLogInterval);
                                        audioLog.textContent = '';
                                        statusDiv.textContent = '녹음된 오디오를 처리중입니다...';

                                        const audioBlob = new Blob(audioChunks, { type: 'audio/webm;codecs=opus' });
                                        console.log(`Total audio size: ${audioBlob.size} bytes`);

                                        if (audioBlob.size === 0) {
                                            throw new Error('녹음된 오디오 데이터가 없습니다.');
                                        }

                                        // 오디오 재생기 설정
                                        const audioURL = URL.createObjectURL(audioBlob);
                                        const audioPlayback = document.getElementById('audioPlayback');
                                        audioPlayback.src = audioURL;
                                        audioPlayback.style.display = 'block';

                                        // Blob을 base64로 변환 및 파일 저장
                                        console.log("Converting audio blob to base64 and saving file...");
                                        const reader = new FileReader();
                                        reader.readAsDataURL(audioBlob);

                                        reader.onloadend = function () {
                                            const base64Audio = reader.result.split(',')[1]; // Remove data URL prefix
                                            const timestamp = new Date().toISOString().replace(/[:.]/g, '-');
                                            const filename = `recorded_${timestamp}.webm`;

                                            // Streamlit 세션 상태 업데이트
                                            window.parent.postMessage({
                                                type: 'streamlit:setComponentValue',
                                                value: JSON.stringify({
                                                    audio: base64Audio,
                                                    filename: filename,
                                                }),
                                            }, '*');

                                            // 추가: 다운로드 링크 제공
                                            const downloadLink = document.createElement('a');
                                            downloadLink.href = audioURL;
                                            downloadLink.download = filename;
                                            downloadLink.textContent = '오디오 파일 다운로드';
                                            document.body.appendChild(downloadLink);
                                        };

                                    } catch (error) {
                                        console.error('Error:', error);
                                        errorDiv.textContent = `오류 발생: ${error.message}`;
                                        statusDiv.textContent = '오류가 발생했습니다. 다시 시도해 주세요.';
                                    }
                                };

                                mediaRecorder.start();
                                isRecording = true;
                                recordButton.textContent = '녹음 중지';
                                recordButton.style.backgroundColor = '#f44336';
                                statusDiv.textContent = '녹음 중...';
                                errorDiv.textContent = '';
                                
                                let startTime = Date.now();
                                audioLogInterval = setInterval(() => {
                                    const duration = Math.floor((Date.now() - startTime) / 1000);
                                    audioLog.textContent = `녹음 시간: ${duration}초`;
                                }, 1000);
                                
                            } else {
                                mediaRecorder.stop();
                                isRecording = false;
                                recordButton.textContent = '녹음 시작';
                                recordButton.style.backgroundColor = '#4CAF50';
                            }
                        } catch (error) {
                            console.error('Error:', error);
                            errorDiv.textContent = `오류 발생: ${error.message}`;
                            statusDiv.textContent = '오류가 발생했습니다.';
                        }
                    }
                </script>
            """

            component_value = st.components.v1.html(
                recording_html,
                height=200
            )
            
            # 컴포넌트 값이 있고 문자열인 경우에만 처리
            if component_value and isinstance(component_value, str):
                try:
                    # JSON 파싱
                    data = json.loads(component_value)
                    audio_base64 = data.get('audio')
                    filename = data.get('filename')
                    
                    if audio_base64 and filename:
                        # base64 디코딩
                        audio_data = base64.b64decode(audio_base64)
                        
                        # 파일 저장
                        filepath = os.path.join("recorded_audio", filename)
                        with open(filepath, "wb") as f:
                            f.write(audio_data)
                        
                        st.success(f"녹음 파일이 저장되었습니다: {filename}")
                        
                        # 오디오 처리
                        result = st.session_state.generator.stt_manager.process_webm_to_text(filepath)
                        
                        if result:
                            st.success("음성 인식이 완료되었습니다!")
                            st.write("인식된 텍스트:")
                            st.write(result)
                            st.session_state.transcribed_text = result
                        else:
                            st.error(f"오류가 발생했습니다: 오디오 처리 실패")
            
                except json.JSONDecodeError:
                    # JSON 파싱 오류 무시 (컴포넌트가 초기화될 때 발생할 수 있음)
                    pass
                except Exception as e:
                    st.error(f"오디오 처리 중 오류가 발생했습니다: {str(e)}")
        
            # 업로드된 파일 처리
            uploaded_file = st.file_uploader("또는 오디오 파일을 업로드하세요", type=['webm', 'wav', 'mp3'])
            
            if uploaded_file is not None:
                st.audio(uploaded_file, format='audio/webm')
                
                if st.button("음성 인식 시작"):
                    try:
                        # 파일 데이터 읽기
                        audio_data = uploaded_file.read()
                        
                        # 오디오 처리
                        result = st.session_state.generator.stt_manager.process_webm_to_text(uploaded_file.name)
                        
                        if result:
                            st.success("음성 인식이 완료되었습니다!")
                            st.write("인식된 텍스트:")
                            st.write(result)
                            st.session_state.transcribed_text = result
                        else:
                            st.error(f"오류가 발생했습니다: 오디오 처리 실패")
                
                    except Exception as e:
                        st.error(f"오디오 처리 중 오류가 발생했습니다: {str(e)}")

            # 소장 생성 버튼
            if st.button("소장 생성", type="primary"):
                if consultation_text:
                    with st.spinner("소장을 생성하고 있습니다..."):
                        try:
                            # 소장 생성
                            complaint = st.session_state.generator.generate_complaint(consultation_text)
                            st.session_state.generated_complaint = complaint
                            
                            # 세션 기록과 MongoDB에 저장
                            st.session_state.session_manager.save_session(consultation_text, complaint)
                            st.session_state.generator.mongo_db.save_conversation(
                                st.session_state.session_manager.get_sessions()[-1].session_id,
                                consultation_text,
                                complaint
                            )
                            
                            st.rerun()
                        except Exception as e:
                            st.error(f"오류가 발생했습니다: {str(e)}")
                else:
                    st.warning("상담 내용을 입력해주세요.")

            # 소장 평가 및 수정 UI 표시
            display_complaint_actions()

if __name__ == "__main__":
    main()