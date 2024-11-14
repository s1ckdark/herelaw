import openai
import json
from typing import Dict, Any
from datetime import datetime
from docx import Document
from dotenv impoert load_dotenv

class DivorceComplaintGenerator:
    def __init__(self, api_key: str):
        self.client=openai.OpenAI(api_key=api_key)
        # JSON 템플릿 파일 로드
        with open('complaint_template.json', 'r', encoding='utf-8') as f:
            self.template = json.load(f)

    def read_dialog_from_docx(self, file_path: str) -> str:
        """docx 파일에서 대화 내용 읽기"""
        doc = Document(file_path)
        dialog_text = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        return '\n'.join(dialog_text)

    def extract_information_from_dialog(self, dialog_text: str) -> Dict[str, Any]:
        """대화에서 필요한 정보를 추출하여 JSON 형식으로 반환"""
        system_prompt = """
        다음 대화에서 이혼 소장 작성에 필요한 정보를 추출하여 JSON 형식으로 반환하세요.
        누락된 정보는 None으로 유지하고, 필요한 형식으로 반환하세요.
        """
        
        # OpenAI API 요청
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": dialog_text}
                ],
                temperature=0
            )
            # 응답 내용 출력하여 확인
            print("API 응답:", response)

            # 응답이 비어 있지 않은지 확인 후 JSON 디코딩
            content = response.choices[0].message.content
            if content:
                return json.loads(content)
            else:
                print("API 응답이 비어 있습니다.")
                return {}

        except json.JSONDecodeError as e:
            print("JSON 디코딩 오류:", e)
            print("응답 내용:", response.choices[0].message)
            return {}

    def create_complaint_document(self, data: Dict[str, Any], output_path: str):
        """JSON 데이터를 사용하여 소장 문서 작성"""
        doc = Document()
        doc.add_heading("이혼 청구의 소", level=1)
        doc.add_paragraph(f"작성일자: {datetime.now().strftime('%Y년 %m월 %d일')}")
        
        # 원고 정보
        doc.add_heading("원고 정보", level=2)
        doc.add_paragraph(f"원고: {data['당사자']['원고']['성명']} (주민등록번호: {data['당사자']['원고']['주민등록번호']})")
        doc.add_paragraph(f"주소: {data['당사자']['원고']['주소']} (우편번호: {data['당사자']['원고']['우편번호']})")

        # 피고 정보
        doc.add_heading("피고 정보", level=2)
        for idx, 피고 in enumerate(data['당사자']['피고'], 1):
            doc.add_paragraph(f"피고 {idx}: {피고['성명']} (주민등록번호: {피고['주민등록번호']})")

        # 청구취지 작성
        doc.add_heading("청구 취지", level=2)
        if data["청구취지"]["이혼청구"]["청구여부"]:
            doc.add_paragraph("1. 원고와 피고는 이혼한다.")
        
        if data["청구취지"]["위자료"]["청구여부"]:
            doc.add_paragraph(f"2. 피고는 원고에게 위자료로 금 {data['청구취지']['위자료']['청구금액']}원을 지급하라.")
        
        # 문서 저장
        doc.save(output_path)
        print(f"소장 문서가 '{output_path}'에 저장되었습니다.")

def main():
    load_dotenv()
    api_key = os.getenv("HERELAW_OPENAI_API_KEY")
    generator = DivorceComplaintGenerator(api_key)

    # 대화 내용이 담긴 docx 파일 경로
    dialog_path = "data.docx"
    output_path = "generated_complaint.docx"

    # 대화 내용 읽기 및 정보 추출
    dialog_text = generator.read_dialog_from_docx(dialog_path)
    extracted_data = generator.extract_information_from_dialog(dialog_text)

    # 추출된 데이터를 기반으로 소장 작성
    generator.create_complaint_document(extracted_data, output_path)

if __name__ == "__main__":
    main()

