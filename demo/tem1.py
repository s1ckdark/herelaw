from typing import Dict, Any
import json
import openai
from datetime import datetime
from docx import Document
from dotenv import load_dotenv
import os

class DivorceComplaintGenerator:
    def __init__(self, api_key: str):
        self.client = openai.OpenAI(api_key = api_key)
        # JSON 템플릿을 파일에서 로드
        with open('complaint_template.json', 'r', encoding='utf-8', errors='ignore') as f:
            self.template = json.load(f, strict=False)

    def read_dialog_from_docx(self, file_path: str) -> str:
        """docx 파일에서 대화 내용을 읽어와 텍스트로 반환"""
        doc = Document(file_path)
        dialog_text = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        return '\n'.join(dialog_text)

    def extract_information_from_dialog(self, dialog_text: str) -> Dict[str, Any]:
        """대화에서 정보를 추출하여 JSON 형식으로 반환"""
        system_prompt = """
        다음 대화에서 이혼 소장 작성에 필요한 정보를 추출하여 JSON 형식으로 반환하세요.
        누락된 정보는 None으로 유지하고, 필요한 형식으로 반환하세요.
        """
        response = self.client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": dialog_text}
            ],
            temperature=0
        )
        return json.loads(response.choices[0].message.content)

    def create_complaint_document(self, data: Dict[str, Any], output_path: str):
        """JSON 데이터를 사용하여 소장 문서를 생성하여 .docx 파일로 저장"""
        doc = Document()

        # 제목 및 기본정보 작성
        doc.add_heading("이혼 청구의 소", level=1)
        doc.add_paragraph(f"작성일자: {datetime.now().strftime('%Y년 %m월 %d일')}")
        
        # 원고 정보
        doc.add_heading("원고 정보", level=2)
        doc.add_paragraph(f"원고: {data['소장']['당사자']['원고']['성명']} (주민등록번호: {data['소장']['당사자']['원고']['주민등록번호']})")
        doc.add_paragraph(f"주소: {data['소장']['당사자']['원고']['주소']} (우편번호: {data['소장']['당사자']['원고']['우편번호']})")

        # 원고 소송대리인 정보
        doc.add_paragraph(f"원고 소송대리인: 변호사 {data['소장']['당사자']['원고대리인']['성명']}")

        # 피고 정보
        doc.add_heading("피고 정보", level=2)
        for idx, 피고 in enumerate(data['소장']['당사자']['피고'], 1):
            doc.add_paragraph(f"피고 {idx}: {피고['성명']} (주민등록번호: {피고['주민등록번호']})")

        # 청구취지 작성
        doc.add_heading("청구 취지", level=2)
        if data["소장"]["청구취지"]["이혼청구"]["청구여부"]:
            doc.add_paragraph("1. 원고와 피고는 이혼한다.")
        
        if data["소장"]["청구취지"]["위자료"]["청구여부"]:
            doc.add_paragraph(f"2. 피고는 원고에게 위자료로 금 {data['소장']['청구취지']['위자료']['청구금액']}원을 지급하라.")
        
        if data["소장"]["청구취지"]["재산분할"]["청구여부"]:
            doc.add_paragraph(f"3. 피고는 원고에게 재산분할로 금 {data['소장']['청구취지']['재산분할']['청구금액']}원을 지급하라.")

        if data["소장"]["청구취지"]["양육비"]["청구여부"]:
            doc.add_paragraph("4. 피고는 미성년 자녀에 대한 양육비로 아래 금액을 지급하라.")
            for 자녀 in data["소장"]["청구취지"]["양육비"]["자녀별내역"]:
                doc.add_paragraph(f"   - 자녀 순번 {자녀['자녀순번']}: 월 {자녀['월지급액']}원, 지급기간: {자녀['지급기간']['시작일']}부터 {자녀['지급기간']['종료일']}까지")

        # 소송비용
        if data["소장"]["청구취지"]["소송비용"]["부담자"]:
            doc.add_paragraph(f"5. 소송 비용은 {data['소장']['청구취지']['소송비용']['부담자']}가 부담한다.")

        # 가집행 선고
        if data["소장"]["청구취지"]["가집행선고"]["신청여부"]:
            대상항목 = ", ".join(data["소장"]["청구취지"]["가집행선고"]["대상항목"])
            doc.add_paragraph(f"6. {대상항목}에 대하여 가집행 선고를 한다.")

        # 청구원인 작성
        doc.add_heading("청구 원인", level=2)
        혼인관계 = data['소장']['청구원인']['당사자관계']['혼인관계']
        doc.add_paragraph(f"혼인 신고일: {혼인관계['혼인신고일']}, 혼인 기간: {혼인관계['혼인기간']}, 자녀 수: {혼인관계['자녀수']}")

        # 이혼 사유
        doc.add_paragraph("이혼 사유:")
        for reason in data['소장']['청구원인']['이혼사유']['사유내용']:
            doc.add_paragraph(f" - 유형: {reason['유형']}, 내용: {reason['구체적내용']}, 발생 시기: {reason['발생시기']['시작']}부터 {reason['발생시기']['종료']}까지")

        # 입증 방법
        doc.add_heading("입증 방법", level=2)
        for 증거 in data["소장"]["입증방법"]["기타증거"]:
            doc.add_paragraph(f" - {증거['증거명']}: {증거['입증취지']}")

        # 첨부서류
        doc.add_heading("첨부 서류", level=2)
        첨부서류 = data["소장"]["첨부서류"]
        doc.add_paragraph(f" - 소송 위임장 제출 여부: {'제출' if 첨부서류['소송위임장']['제출여부'] else '미제출'}")
        doc.add_paragraph(f" - 송달료 납부 증명 제출 여부: {'제출' if 첨부서류['송달료납부증명']['제출여부'] else '미제출'}, 납부 금액: {첨부서류['송달료납부증명'].get('납부금액', '없음')}")
        doc.add_paragraph(f" - 인지 첨부 증명 제출 여부: {'제출' if 첨부서류['인지첨부증명']['제출여부'] else '미제출'}, 인지 액: {첨부서류['인지첨부증명'].get('인지액', '없음')}")

        # 파일 저장
        doc.save(output_path)
        print(f"소장 문서가 '{output_path}'에 저장되었습니다.")

def main():
    api_key = os.getenv("HERELAW_OPENAI_API_KEY")
    generator = DivorceComplaintGenerator(api_key)

    # 대화 내용이 담긴 docx 파일 경로
    dialog_path = "data.docx"
    output_path = "generated_complaint.docx"

    # 대화 내용 읽기 및 정보 추출
    dialog_text = generator.read_dialog_from_docx(dialog_path)
    extracted_data = generator.extract_information_from_dialog(dialog_text)

    # 추출된 데이터를 기반으로 소장 작성
    generator.create_complaint_document(extracted_data, output_path)

if __name__ == "__main__":
    main()

