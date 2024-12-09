from flask import Flask, request, jsonify
from openai import OpenAI

client = OpenAI()
import os
from dotenv import load_dotenv
from flask_cors import CORS
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain_community.llms import OpenAI
from pymongo import MongoClient
from datetime import datetime, timedelta
import uuid
import hashlib
from sklearn.preprocessing import StandardScaler
import numpy as np
from docx import Document
from langsmith import Client
from typing import List, Dict, Optional
import json
from moviepy import AudioFileClip
import jwt
from werkzeug.security import generate_password_hash, check_password_hash
from functools import wraps
from bson import ObjectId
import speech_recognition as sr
import os
from werkzeug.utils import secure_filename
import subprocess
from io import BytesIO
from flask import send_file
from bson.errors import InvalidId
from flask_sock import Sock
import tempfile
import wave
import io

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": ["http://herelaw.nomadseoul.com", "http://localhost:3000"], "supports_credentials": True}})
app.config['JWT_SECRET_KEY'] = os.getenv('JWT_SECRET_KEY', 'your-secret-key')  # 실제 배포 시에는 반드시 환경 변수로 설정해야 합니다
app.config['UPLOAD_FOLDER'] = './uploads'

sock = Sock(app)

# Load environment variables
load_dotenv()

class MongoDBManager:
    def __init__(self, uri):
        if not uri:
            raise ValueError("MongoDB URI is required")
        self.client = MongoClient(uri)
        self.db = self.client[os.getenv('MONGODB_DB', 'herelaw')]
        self.conversations = self.db[os.getenv('MONGODB_COLLECTION', 'divorce_complaint')]
        self.documents = self.db['document_chunks']
        self.feedback = self.db['feedback']
        self.users = self.db['users']
        self.sessions = self.db['sessions']
        self.logs = self.db['logs']  # Add logs collection

        print(f"MongoDB 연결 정보:")
        print(f"Database: {self.db.name}")
        print(f"Collections: {self.db.list_collection_names()}")

        # Create indexes
        self.documents.create_index([("chunk_hash", 1)], unique=True)
        self.feedback.create_index([("session_id", 1)])
        self.users.create_index([("username", 1)], unique=True)
        self.users.create_index([("email", 1)], unique=True)
        self.sessions.create_index([("user_id", 1)])
        self.logs.create_index([("user_id", 1)])  # Add logs index

    def create_user(self, username: str, password: str, email: str) -> Optional[str]:
        """새 사용자를 생성합니다."""
        try:
            print(f"사용자 생성 시도: {username}")
            user_id = str(uuid.uuid4())
            result = self.users.insert_one({
                "user_id": user_id,
                "username": username,
                "email": email,
                "password": generate_password_hash(password),
                "created_at": datetime.utcnow(),
                "last_login": None,
                "role": "user",  # Default role
                "level": 1,      # Default level
                "status": "active"
            })
            print(f"사용자 생성 결과: {result.inserted_id}")
            return user_id
        except Exception as e:
            print(f"사용자 생성 중 오류 발생: {str(e)}")
            return None

    def get_user_by_username(self, username: str) -> Optional[dict]:
        """사용자명으로 사용자를 조회합니다."""
        return self.users.find_one({"username": username})

    def get_user_by_id(self, user_id: str) -> Optional[dict]:
        """사용자 ID로 사용자를 조회합니다."""
        return self.users.find_one({"user_id": user_id})

    def verify_user(self, username: str, password: str) -> Optional[dict]:
        """사용자 인증을 수행합니다."""
        user = self.get_user_by_username(username)
        if user and check_password_hash(user["password"], password):
            self.users.update_one(
                {"user_id": user["user_id"]},
                {"$set": {"last_login": datetime.utcnow()}}
            )
            return user
        return None

    def save_session(self, user_id: str, consultation_text: str, generated_content: str) -> str:
        """사용자 세션을 저장합니다."""
        session_id = str(uuid.uuid4())
        self.sessions.insert_one({
            "session_id": session_id,
            "user_id": user_id,
            "consultation_text": consultation_text,
            "generated_content": generated_content,
            "created_at": datetime.utcnow(),
            "rating": None,
            "feedback": None
        })
        return session_id

    def get_user_sessions(self, user_id: str) -> List[dict]:
        """사용자의 모든 세션을 조회합니다."""
        return list(self.sessions.find(
            {"user_id": user_id},
            {"_id": 0}
        ).sort("created_at", -1))

    def update_session(self, session_id: str, rating: Optional[int] = None, 
                      feedback: Optional[str] = None) -> bool:
        """세션을 업데이트합니다."""
        update_data = {}
        if rating is not None:
            update_data["rating"] = rating
        if feedback is not None:
            update_data["feedback"] = feedback

        if update_data:
            result = self.sessions.update_one(
                {"session_id": session_id},
                {"$set": update_data}
            )
            return result.modified_count > 0
        return False

    def save_chunk(self, content: str, doc_type: str, embedding: List[float]):
        chunk_hash = hashlib.md5(content.encode()).hexdigest()
        if not self.documents.find_one({"chunk_hash": chunk_hash}):
            self.documents.insert_one({
                "content": content,
                "doc_type": doc_type,
                "chunk_hash": chunk_hash,
                "embedding": embedding,
                "created_at": datetime.now()
            })

    def save_conversation(self, session_id: str, user_input: str, generated_content: Dict):
        conversation = {
            "session_id": session_id,
            "created_at": datetime.now(),
            "user_input": user_input,
            "generated_content": generated_content
        }
        self.conversations.insert_one(conversation)
        return conversation

    def get_similar_chunks(self, query_embedding: List[float], doc_type: str, k: int = 3):
        """벡터 유사도 검색을 수행합니다."""
        pipeline = [
            {
                "$search": {
                    "index": "vector_index",
                    "knnBeta": {
                        "vector": query_embedding,
                        "path": "embedding",
                        "k": k
                    }
                }
            },
            {
                "$match": {
                    "doc_type": doc_type
                }
            }
        ]

        results = list(self.documents.aggregate(pipeline))
        return results

    def save_feedback(self, feedback_doc):
        """피드백을 저장합니다."""
        try:
            # 입력값 검증
            if not isinstance(feedback_doc.get('rating'), (int, float)):
                raise ValueError("rating must be a number")

            result = self.feedback.insert_one(feedback_doc)

            if result.inserted_id:
                print(f"피드백 저장 성공: {result.inserted_id}")
                return result.inserted_id
            else:
                raise Exception("피드백 저장 실패")

        except Exception as e:
            print(f"피드백 저장 중 오류: {str(e)}")
            raise e

    def get_feedback_statistics(self):
        pipeline = [
            {
                "$group": {
                    "_id": None,
                    "avg_rating": {"$avg": "$rating"},
                    "count": {"$sum": 1},
                    "ratings": {"$push": "$rating"}
                }
            }
        ]

        stats = list(self.feedback.aggregate(pipeline))
        if stats:
            return {
                "average_rating": stats[0]["avg_rating"],
                "total_feedback": stats[0]["count"],
                "rating_distribution": np.histogram(stats[0]["ratings"], bins=5)[0].tolist()
            }
        return None

    def save_log(self, log_data):
        """로그를 저장합니다."""
        try:
            result = self.logs.insert_one(log_data)
            if result.inserted_id:
                print(f"로그 저장 성공: {result.inserted_id}")
                return result.inserted_id
            else:
                raise Exception("로그 저��� 실패")

        except Exception as e:
            print(f"로그 저장 중 오류: {str(e)}")
            raise e
class ReinforcementLearner:
    def __init__(self, mongo_db):
        self.mongo_db = mongo_db
        self.scaler = StandardScaler()
        self.min_feedback_samples = 10

    def extract_features(self, complaint: str) -> List[float]:
        """소장에서 특징 추출"""
        features = []

        # 1. 문서 길이
        features.append(len(complaint))

        # 2. 섹션 수
        sections = [s for s in complaint.split('\n\n') if s.strip()]
        features.append(len(sections))

        # 3. 법률 용어 사용 빈도
        legal_terms = ['청구취지', '청구원인', '입증방법', '첨부서류']
        for term in legal_terms:
            features.append(complaint.count(term))

        # 4. 문장의 평균 길이
        sentences = [s.strip() for s in complaint.split('.') if s.strip()]
        avg_sentence_length = sum(len(s) for s in sentences) / len(sentences) if sentences else 0
        features.append(avg_sentence_length)

        return features

    def calculate_reward(self, user_rating: float, complaint_length: int) -> float:
        """보상 계산"""
        # 기본 보상: 사용자 평가 (1-5점)
        base_reward = user_rating

        # 길이 보상: 적정 길이(1000-3000자)일 때 최대
        length_reward = 0
        if 1000 <= complaint_length <= 3000:
            length_reward = 1
        elif complaint_length < 1000:
            length_reward = complaint_length / 1000
        else:
            length_reward = 3000 / complaint_length

        # 최종 보상 = 기본 보상 * 0.7 + 길이 보상 * 0.3
        final_reward = base_reward * 0.7 + length_reward * 0.3

        return final_reward

    def get_best_practices(self) -> Dict:
        """높은 평가를 받은 소장들의 특징을 분석"""
        try:
            # 피드백 데이터 가져오기
            feedback_data = list(self.mongo_db.feedback.find(
                {"rating": {"$gte": 4}}  # 4점 이상 평가받은 소장만 선택
            ).sort("rating", -1).limit(50))  # 상위 50개

            if len(feedback_data) < self.min_feedback_samples:
                return None

            # 높은 평가를 받은 소장들만 선택 (평점 4점 이상)
            successful_complaints = [
                data for data in feedback_data 
                if data.get('rating', 0) >= 4
            ]

            if not successful_complaints:
                return None

            # 1. 평균 문서 길이 계산
            avg_length = sum(len(c['complaint']) for c in successful_complaints) / len(successful_complaints)

            # 2. 자주 사용된 문구 추출
            common_phrases = self._extract_common_phrases(successful_complaints)

            # 3. 섹션 구조 분석
            section_patterns = self._analyze_section_patterns(successful_complaints)

            # 4. 성공적인 특징 분석
            successful_features = self._analyze_successful_features(successful_complaints)

            return {
                'avg_length': avg_length,
                'common_phrases': common_phrases,
                'section_patterns': section_patterns,
                'successful_features': successful_features
            }

        except Exception as e:
            print(f"분석 중 오류 발생: {str(e)}")
            return None

    def _extract_common_phrases(self, feedback_data: List[Dict]) -> List[str]:
        """성공적인 소장들에서 자주 사용된 문구 추출"""
        # 여기에 문구 추출 로직 구현
        all_phrases = []
        for data in feedback_data:
            complaint = data["complaint"]
            # 문장 단위로 분리
            sentences = complaint.split('.')
            # 3단어 이상의 문구만 추출
            phrases = [s.strip() for s in sentences if len(s.strip().split()) >= 3]
            all_phrases.extend(phrases)

        # 빈도수 기준으로 상위 10개 문구 반환
        from collections import Counter
        phrase_counter = Counter(all_phrases)
        return [phrase for phrase, count in phrase_counter.most_common(10)]

    def _analyze_section_patterns(self, feedback_data: List[Dict]) -> Dict:
        """성공적인 소장들의 섹션 구조 분석"""
        sections = ['청구취지', '청구원인', '입증방법', '첨부서류']
        section_stats = {
            'section_order': [],
            'section_lengths': {},
            'common_transitions': {}
        }

        for data in feedback_data:
            complaint = data["complaint"]
            # 섹션 순서 분석
            found_sections = []
            for section in sections:
                if section in complaint:
                    found_sections.append(section)

            if found_sections:
                section_stats['section_order'].append(found_sections)

        # 가장 흔한 섹션 순서 찾기
        from collections import Counter
        order_counter = Counter(tuple(order) for order in section_stats['section_order'])
        section_stats['most_common_order'] = list(order_counter.most_common(1)[0][0])

        return section_stats

    def _analyze_successful_features(self, feedback_data: List[Dict]) -> Dict:
        """성공적인 소장들의 특징 분석"""
        features = {
            'avg_length': 0,
            'common_legal_terms': {},
            'section_coverage': {},
            'style_patterns': {}
        }

        legal_terms = ['원고', '피고', '위자료', '재산분할', '양육권', '가집행']
        total_complaints = len(feedback_data)

        for data in feedback_data:
            complaint = data["complaint"]

            # 길이 분석
            features['avg_length'] += len(complaint)

            # 법률 용어 ��용 빈도
            for term in legal_terms:
                if term not in features['common_legal_terms']:
                    features['common_legal_terms'][term] = 0
                features['common_legal_terms'][term] += complaint.count(term)

        # 평균 계산
        features['avg_length'] /= total_complaints
        for term in features['common_legal_terms']:
            features['common_legal_terms'][term] /= total_complaints

        return features

class DivorceComplaintGenerator:
    def __init__(self):
        # 환 변수에서 값 로드
        api_key = os.getenv("OPENAI_API_KEY")
        mongo_uri = os.getenv("MONGO_URI")
        endpoint = os.getenv("LANGCHAIN_ENDPOINT", "https://api.smith.langchain.com")
        langchain_api_key = os.getenv("LANGCHAIN_API_KEY")
        self.project_name = os.getenv("LANGCHAIN_PROJECT", "default_project")  # 기본값 설정

        # LangSmith 설정
        if langchain_api_key:  # API 키가 있을 때만 LangSmith 초기화
            print("LangSmith 클라이언트 초기화")
            self.langsmith_client = Client(
                api_url=endpoint,  # base_url로 변경
                api_key=langchain_api_key,
            )
        else:
            self.langsmith_client = None
            print("LangSmith 클라이언트가 초기화되지 않았습니다. 트레이싱이 비활성화됩니다.")

        # OpenAI Embeddings 초기화
        self.embeddings = OpenAIEmbeddings(api_key=api_key)

        # MongoDB 연결
        self.mongo_db = MongoDBManager(mongo_uri)
        self.session_manager = SessionManager()
        self.rl_learner = ReinforcementLearner(self.mongo_db)
        self.user_manager = UserManager(self.mongo_db)  # UserManager 인스턴스 생성

    def generate_complaint(self, consultation_text: str) -> dict:
        """소장 생성"""
        run_id = None

        if self.langsmith_client:  # LangSmith가 설정된 경우 실행 생성
            # 실행 생성 (name 및 run_type 포함)
            run_id = self.langsmith_client.create_run(
                name="generate_complaint",  # 실행 이름
                run_type="tool",            # 실행 유형 (예: 'tool', 'chain', 'llm' 등)
                project_name=self.project_name,  # 프로젝트 이름
                inputs={"consultation_text": consultation_text}  # 입력 데이터
            )

        try:
            # 내부 로직 실행
            result = self._generate_complaint_internal(consultation_text)

            # 실행 결과 업데이트
            if self.langsmith_client and run_id:
                self.langsmith_client.update_run(
                    run_id=run_id,
                    outputs=result,
                    status="completed"
                )
            return result
        except Exception as e:
            # 실행 오류 처리
            if self.langsmith_client and run_id:
                self.langsmith_client.update_run(
                    run_id=run_id,
                    error=str(e),
                    status="failed"
                )
            raise e


    def _generate_complaint_internal(self, consultation_text: str) -> dict:
        """실제 소장 생성 로직 (임의로 대체 가능)"""
        claim_chunks = ["Claim data placeholder"]
        relief_chunks = ["Relief data placeholder"]

        return self._generate_with_gpt(consultation_text, claim_chunks, relief_chunks)

    def _generate_with_gpt(self, consultation_text: str, claim_chunks: List[str], relief_chunks: List[str]) -> dict:
        """GPT로 소장 생성 (피드백 학습 적용)"""
        # 성공적인 소장의 특징 가져오기
        best_practices = self.rl_learner.get_best_practices()

        if best_practices:
            # 피드백 기반 프롬프트 최적화
            quality_guidelines = f"""
            다음 기준을 충족하는 고품질 소장을 작성해주세요:
            1. 적정 문서 길이: {int(best_practices['avg_length'])} 자 내외
            2. 다음 문구들을 적절히 활용하세요:
               {', '.join(best_practices['common_phrases'][:5])}
            3. 성공적인 소장의 섹션 구조를 따르세요:
               {best_practices['section_patterns']}
            4. 다음 특징들을 반영하세요:
               {best_practices['successful_features']}
            """
        else:
            # 기본 가이드라인 사용
            quality_guidelines = """
            다음 기준을 충족하는 고품질 소장을 작성해주세요:
            1. 모든 필수 섹션을 포함할 것
            2. 구체적이고 명확한 법률 용어 사용
            3. 논리적인 구조와 흐름
            4. 적절한 길이와 상세도
            """

        prompt = (
            quality_guidelines + "\n\n" +
            "다음 상담 내용과 참고 문서를 바탕으로 이혼 소장을 작성해주세요.\n\n" +
            f"[상담 내용]\n{consultation_text}\n\n" +
            f"[청구취지 참고문서]\n{' '.join(claim_chunks)}\n\n" +
            f"[청구원인 참고문서]\n{' '.join(relief_chunks)}\n\n"
        )

        # GPT 모델에 컨텍스트 추가
        messages = [
            {
                "role": "system",
                "content": "당신은 전문 법률 문서 작성 시스템입니다. 과거의 성공적인 소장 작성 경험을 바탕으로 최적화된 이혼 소장을 작성합니다."
            }
        ]

        # 성공적인 예시 추가
        if best_practices and 'example_complaints' in best_practices:
            messages.append({
                "role": "assistant",
                "content": f"다음은 높은 평가를 받은 소장의 예시입니다:\n{best_practices['example_complaints'][0]}"
            })

        messages.append({
            "role": "user",
            "content": prompt
        })

        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            temperature=0.3
        )

        return response.choices[0].message.content

class UserManager:
    def __init__(self, mongodb_manager):
        self.mongodb = mongodb_manager
        self.users_collection = self.mongodb.users

    def create_user(self, username: str, password: str, email: str) -> Optional[str]:
        """새 사용자를 생성합니다."""
        try:
            # 중복 확인
            if self.users_collection.find_one({"username": username}):
                raise ValueError("이미 존재하는 사용자명입니다.")
            if self.users_collection.find_one({"email": email}):
                raise ValueError("이미 등록된 이메일입니다.")

            user_id = str(uuid.uuid4())
            user_doc = {
                "user_id": user_id,
                "username": username,
                "email": email,
                "password": generate_password_hash(password),
                "created_at": datetime.utcnow(),
                "last_login": None,
                "role": "user",
                "status": "active",
                "level": 1,
                "settings": {
                    "notification_enabled": True,
                    "theme": "light"
                }
            }

            result = self.users_collection.insert_one(user_doc)
            if result.inserted_id:
                return user_id
            return None

        except Exception as e:
            print(f"사용자 생성 중 오류: {str(e)}")
            raise

    def get_user(self, user_id: str) -> Optional[dict]:
        """사용자 정보를 조회합니다."""
        try:
            user = self.users_collection.find_one({"user_id": user_id})
            if user:
                user.pop('password', None)  # 비밀번호 필드 제거
                return user
            return None
        except Exception as e:
            print(f"사용자 조회 중 오류: {str(e)}")
            return None

    def update_user(self, user_id: str, updates: dict) -> bool:
        """사용자 정보를 업데이트합니다."""
        try:
            # 보안을 위해 업데이트 불가능한 필드 제거
            updates.pop('user_id', None)
            updates.pop('password', None)
            updates.pop('role', None)

            result = self.users_collection.update_one(
                {"user_id": user_id},
                {"$set": updates}
            )
            return result.modified_count > 0
        except Exception as e:
            print(f"사용자 업데이트 중 오류: {str(e)}")
            return False

    def change_password(self, user_id: str, current_password: str, new_password: str) -> bool:
        """사용자 비밀번호를 변경합니다."""
        try:
            user = self.users_collection.find_one({"user_id": user_id})
            if not user:
                return False

            # 현재 비밀번호 확인
            if not check_password_hash(user['password'], current_password):
                return False

            # 새 비밀번호로 업데이트
            result = self.users_collection.update_one(
                {"user_id": user_id},
                {"$set": {"password": generate_password_hash(new_password)}}
            )
            return result.modified_count > 0
        except Exception as e:
            print(f"비밀번호 변경 중 오류: {str(e)}")
            return False

    def delete_user(self, user_id: str) -> bool:
        """사용자 계정을 비활성화합니다."""
        try:
            result = self.users_collection.update_one(
                {"user_id": user_id},
                {"$set": {"status": "inactive", "deactivated_at": datetime.utcnow()}}
            )
            return result.modified_count > 0
        except Exception as e:
            print(f"사용자 삭제 중 오류: {str(e)}")
            return False

    def verify_user(self, username: str, password: str) -> Optional[dict]:
        """사용자 인증을 수행합니다."""
        try:
            user = self.users_collection.find_one({"username": username})
            if user and check_password_hash(user['password'], password):
                # 마지막 로그인 시간 업데이트
                self.users_collection.update_one(
                    {"user_id": user["user_id"]},
                    {"$set": {"last_login": datetime.utcnow()}}
                )
                user.pop('password', None)  # 비밀번호 필드 제거
                return user
            return None
        except Exception as e:
            print(f"사용자 인증 중 오류: {str(e)}")
            return None

    def get_user_stats(self, user_id: str) -> dict:
        """사용자 통계 정보를 조회합니다."""
        try:
            # 세션 수
            session_count = self.mongodb.sessions.count_documents({"user_id": user_id})

            # 평균 평점
            pipeline = [
                {"$match": {"user_id": user_id}},
                {"$group": {
                    "_id": None,
                    "avg_rating": {"$avg": "$rating"},
                    "total_sessions": {"$sum": 1}
                }}
            ]
            stats = list(self.mongodb.sessions.aggregate(pipeline))

            return {
                "total_sessions": session_count,
                "average_rating": stats[0]["avg_rating"] if stats else 0,
                "last_login": self.get_user(user_id).get("last_login")
            }
        except Exception as e:
            print(f"사용자 통계 조회 중 오류: {str(e)}")
            return {
                "total_sessions": 0,
                "average_rating": 0,
                "last_login": None
            }

    def update_user_level(self, user_id: str) -> bool:
        """사용자 레벨을 업데이트합니다."""
        try:
            stats = self.get_user_stats(user_id)

            # 레벨 계산 로직 (예시)
            new_level = 1
            if stats["total_sessions"] >= 50:
                new_level = 5
            elif stats["total_sessions"] >= 30:
                new_level = 4
            elif stats["total_sessions"] >= 20:
                new_level = 3
            elif stats["total_sessions"] >= 10:
                new_level = 2

            if new_level > 1:
                result = self.users_collection.update_one(
                    {"user_id": user_id},
                    {"$set": {"level": new_level}}
                )
                return result.modified_count > 0
            return False

        except Exception as e:
            print(f"사용자 레벨 업데이트 중 오류: {str(e)}")
            return False
class SessionManager:
    def __init__(self):
        self.db = mongodb_manager.db
        self.sessions_collection = self.db['sessions']

    def create_session(self, user_id: str, consultation_text: str, generated_content: dict = None) -> str:
        """새 세션을 생성합니다."""
        session_id = str(uuid.uuid4())
        session = {
            'session_id': session_id,
            'user_id': user_id,
            'consultation_text': consultation_text,
            'generated_content': generated_content,
            'created_at': datetime.utcnow(),
            'status': 'active'
        }
        self.sessions_collection.insert_one(session)
        return session_id

    def get_session(self, session_id: str, user_id: str) -> Optional[dict]:
        """세션을 조회합니다."""
        return self.sessions_collection.find_one({
            'session_id': session_id,
            'user_id': user_id
        })

    def update_session(self, session_id: str, user_id: str, updates: dict) -> bool:
        """세션을 업데이트합니다."""
        result = self.sessions_collection.update_one(
            {'session_id': session_id, 'user_id': user_id},
            {'$set': updates}
        )
        return result.modified_count > 0

    def get_user_sessions(self, user_id: str) -> List[dict]:
        """사용자의 모든 세션을 조회합니다."""
        return list(self.sessions_collection.find(
            {'user_id': user_id},
            {'_id': 0}
        ).sort('created_at', -1))

class JWTManager:
    def __init__(self, secret_key):
        self.secret_key = secret_key

    def generate_token(self, user_id: str) -> str:
        """사용자 ID로 JWT 토큰을 생성합니다."""
        payload = {
            'user_id': user_id,
            'exp': datetime.utcnow() + timedelta(days=1)  # 토큰 만료 시간: 1일
        }
        return jwt.encode(payload, self.secret_key, algorithm='HS256')

    def verify_token(self, token: str) -> Optional[dict]:
        """JWT 토큰을 검증하고 페이로드를 반환합니다."""
        try:
            payload = jwt.decode(token, self.secret_key, algorithms=['HS256'])
            return payload
        except jwt.ExpiredSignatureError:
            return None
        except jwt.InvalidTokenError:
            return None

# JWT 관리자 초기화
jwt_manager = JWTManager(app.config['JWT_SECRET_KEY'])

# MongoDB 매니저 초기화
mongo_uri = os.getenv('MONGO_URI')
if not mongo_uri:
    raise ValueError("MONGO_URI environment variable is not set")
mongodb_manager = MongoDBManager(mongo_uri)

def jwt_required():
    def decorator(func):
        @wraps(func)
        def decorated_function(*args, **kwargs):
            auth_header = request.headers.get('Authorization')
            if not auth_header or not auth_header.startswith('Bearer '):
                return jsonify({'error': '인증이 필요합니다.'}), 401

            token = auth_header.split(' ')[1]
            payload = jwt_manager.verify_token(token)
            if not payload:
                return jsonify({'error': '유효하지 않은 토큰입니다.'}), 401

            request.user_id = payload['user_id']
            return func(*args, **kwargs)
        return decorated_function
    return decorator

# Admin required decorator
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        token = request.headers.get('Authorization')
        if not token:
            return jsonify({'message': 'Token is missing'}), 401

        try:
            payload = jwt_manager.verify_token(token.split(' ')[1])
            user = mongodb_manager.users.find_one({'_id': ObjectId(payload['user_id'])})
            if not user or user.get('role') != 'admin':
                return jsonify({'message': 'Admin privileges required'}), 403
        except:
            return jsonify({'message': 'Invalid token'}), 401

        return f(*args, **kwargs)
    return decorated_function

# Admin routes
@app.route('/api/admin/users', methods=['GET'])
@admin_required
def get_users():
    users = list(mongodb_manager.users.find({}, {'password': 0}))
    for user in users:
        user['_id'] = str(user['_id'])
    return jsonify(users)

@app.route('/api/admin/users/<user_id>', methods=['PUT'])
@admin_required
def update_user(user_id):
    try:
        data = request.json
        allowed_fields = ['email', 'role', 'status', 'level']
        update_data = {k: v for k, v in data.items() if k in allowed_fields}

        if not update_data:
            return jsonify({'message': 'No valid fields to update'}), 400

        result = mongodb_manager.users.update_one(
            {'_id': ObjectId(user_id)},
            {'$set': update_data}
        )

        if result.modified_count:
            return jsonify({'message': 'User updated successfully'})
        return jsonify({'message': 'User not found'}), 404

    except InvalidId:
        return jsonify({'message': 'Invalid user ID'}), 400

@app.route('/api/admin/users/<user_id>/reset-password', methods=['POST'])
@admin_required
def reset_user_password(user_id):
    try:
        new_password = request.json.get('new_password')
        if not new_password:
            return jsonify({'message': 'New password is required'}), 400

        hashed_password = generate_password_hash(new_password)
        result = mongodb_manager.users.update_one(
            {'_id': ObjectId(user_id)},
            {'$set': {'password': hashed_password}}
        )

        if result.modified_count:
            return jsonify({'message': 'Password reset successfully'})
        return jsonify({'message': 'User not found'}), 404

    except InvalidId:
        return jsonify({'message': 'Invalid user ID'}), 400

@app.route('/api/admin/users/<user_id>/logs', methods=['GET'])
@admin_required
def get_user_logs(user_id):
    try:
        # Get logs from MongoDB (assuming we store logs in a 'logs' collection)
        logs = list(mongodb_manager.db.logs.find({'user_id': user_id}))
        for log in logs:
            log['_id'] = str(log['_id'])
        return jsonify(logs)
    except InvalidId:
        return jsonify({'message': 'Invalid user ID'}), 400

# Modify the existing register route to include user level
@app.route('/api/register', methods=['POST'])
def register():
    """새 사용자를 등록합니다."""
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        email = data.get('email', '')  # email은 선택적으로 처리

        print(f"회원가입 요청 - 사용자명: {username}")

        if not all([username, password]):
            return jsonify({"error": "사용자명과 비밀번호는 필수입니다."}), 400

        # 이미 존재하는 사용자인지 확인
        existing_user = mongodb_manager.get_user_by_username(username)
        if existing_user:
            print(f"이미 존재하는 사용자: {username}")
            return jsonify({"error": "이미 존재하는 사용자명입니다."}), 400

        # 사용자 생성
        user_id = mongodb_manager.create_user(username, password, email)
        if user_id:
            print(f"사용자 생성 성공 - ID: {user_id}")
            # JWT 토큰 생성
            token = jwt_manager.generate_token(user_id)
            return jsonify({
                "message": "회원가입이 완료되었습니다.",
                "token": token,
                "user_id": user_id,
                "username": username
            })
        else:
            print("사용자 생성 실패")
            return jsonify({"error": "사용자 생성에 실패했습니다."}), 500
    except Exception as e:
        print(f"회원가입 중 오류 발생: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/login', methods=['POST'])
def login():
    """사용자 로그인을 처리합니다."""
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')

        if not all([username, password]):
            return jsonify({"error": "사용자명과 비밀번호를 입력해세요."}), 400

        # 사용자 인증
        user = mongodb_manager.verify_user(username, password)
        if user:
            # JWT 토큰 생성
            token = jwt_manager.generate_token(user['user_id'])
            return jsonify({
                "message": "로그인 성공",
                "token": token,
                "user_id": user['user_id'],
                "username": user['username']
            })
        else:
            return jsonify({"error": "잘못된 사용자명 또는 비밀번호입니다."}), 401
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/sessions', methods=['GET'])
@jwt_required()
def get_sessions():
    """사용자의 세션 목록을 반환합니다."""
    try:
        sessions = mongodb_manager.get_user_sessions(request.user_id)
        return jsonify(sessions)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/sessions', methods=['POST'])
@jwt_required()
def create_session():
    """새 세션을 생성합니다."""
    try:
        data = request.get_json()
        consultation_text = data.get('consultation_text')
        generated_content = data.get('generated_content')

        if not all([consultation_text, generated_content]):
            return jsonify({"error": "필 필드가 누락되었습니다."}), 400

        session_id = mongodb_manager.save_session(
            request.user_id,
            consultation_text,
            generated_content
        )

        return jsonify({
            "message": "세션이 생성되었습니다.",
            "session_id": session_id
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/sessions/<session_id>', methods=['PUT'])
@jwt_required()
def update_session(session_id):
    """세션을 업데이트합니다."""
    try:
        data = request.get_json()
        rating = data.get('rating')
        feedback = data.get('feedback')

        if mongodb_manager.update_session(session_id, rating, feedback):
            return jsonify({"message": "세션이 업데이트되었습니다."})
        else:
            return jsonify({"error": "세션 업데이트에 실패했습니다."}), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/sessions/<session_id>', methods=['GET'])
@jwt_required()
def get_session_details(session_id):
    """특정 세션의 상세 정보를 반환합니다."""
    try:
        # 세션 ID 유효성 검사
        if not session_id or session_id == 'undefined':
            return jsonify({"error": "유효하지 않은 세션 ID입니다."}), 400

        # 세션 조회 쿼리 준비
        query = {
            'user_id': request.user_id
        }

        # ObjectId 형식인지 확인
        try:
            # ObjectId로 변환 시도
            object_id = ObjectId(session_id)
            query['_id'] = object_id
        except (InvalidId, TypeError):
            # ObjectId가 아니면 session_id로 검색
            query['session_id'] = session_id

        # 세션 조회
        session = mongodb_manager.sessions.find_one(query)

        if not session:
            return jsonify({"error": "세션을 찾을 수 없습니다."}), 404

        # 날짜 처리 함수
        def parse_date(date_obj):
            if isinstance(date_obj, dict) and '$date' in date_obj:
                return date_obj['$date']
            elif isinstance(date_obj, datetime):
                return date_obj
            else:
                return datetime.now()

        # 세션 데이터 준비 (MongoDB ObjectId를 문자열로 변환)
        session_data = {
            '_id': str(session['_id']),
            'session_id': session.get('session_id', str(session['_id'])),
            'created_at': parse_date(session.get('created_at', datetime.now())),
            'summary': session.get('summary', ''),
            'consultation_text': session.get('consultation_text', []),
            'complaint': session.get('generated_content', ''),
            'title': session.get('title', ''),
            'rating': session.get('rating')
        }

        return jsonify(session_data), 200

    except Exception as e:
        print(f"세션 상세 정보 조회 중 오류: {str(e)}")
        return jsonify({"error": f"세션 정보를 불러오는 중 오류가 발생했습니다: {str(e)}"}), 500

@app.route('/api/generate-complaint', methods=['POST'])
@jwt_required()
def generate_complaint():

    try:
        # 요청 데이터 검증
        complaint_generator = DivorceComplaintGenerator()
        data = request.get_json()
        if not data:
            return jsonify({"error": "요청 데이터가 없습니다."}), 400

        user_input = data.get('user_input')
        if not user_input:
            return jsonify({"error": "사용자 입력이 없습니다."}), 400

        # 현재 사용자 정보 가져오기
        current_user = request.user_id

        try:
            # DivorceComplaintGenerator를 사용하여 소장 생성
            generated_complaint = complaint_generator.generate_complaint(user_input)

            try:
                # 세션 저장
                session_id = mongodb_manager.save_session(
                    request.user_id,
                    user_input,
                    {
                        "complaint": generated_complaint
                    }
                )
            except Exception as db_error:
                print(f"데이터베이스 저장 오류: {str(db_error)}")
                # 데이터베이스 오류가 발생해도 사용자에게는 생성된 응답을 반환

            return jsonify({
                "complaint": generated_complaint,
                "session_id": session_id
            }), 200

        except Exception as e:
            print(f"소장 생성 오류: {str(e)}")
            return jsonify({"error": f"소장 생성 중 오류가 발생했습니다: {str(e)}"}), 500

    except Exception as e:
        print(f"서버 오류: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/rating', methods=['POST'])
@jwt_required()
def save_feedback():
    """피드백을 저장하는 엔드포인트"""
    try:
        data = request.get_json()

        # 필수 필드 검증
        required_fields = ['session_id', 'complaint', 'rating']
        if not all(field in data for field in required_fields):
            missing_fields = [field for field in required_fields if field not in data]
            return jsonify({
                "error": f"필수 필드가 누락되었습니다: {', '.join(missing_fields)}"
            }), 400

        session_id = data['session_id']
        complaint = data['complaint']
        rating = float(data['rating'])

        print(f"피드백 저장 요청 받음: session_id={session_id}, rating={rating}")

        # 피드백 문서 생성
        feedback_doc = {
            "session_id": session_id,
            "complaint": complaint,
            "rating": rating,
            "created_at": datetime.utcnow()
        }

        # MongoDB에 피드백 저장
        feedback_id = mongodb_manager.save_feedback(feedback_doc)

        # 세션 업데이트
        mongodb_manager.update_session(
            session_id=session_id,
            rating=rating,
        )

        print(f"피드백 저장 성공: feedback_id={feedback_id}")

        return jsonify({
            "message": "피드백이 성공적으로 저장되었습니다",
            "feedback_id": str(feedback_id)
        })

    except ValueError as ve:
        print(f"값 류 발생: {str(ve)}")
        return jsonify({"error": f"잘못된 입력값: {str(ve)}"}), 400
    except Exception as e:
        print(f"피드백 저장 중 오류 발생: {str(e)}")
        return jsonify({"error": f"피드백 저장 중 오류가 발생했습니다: {str(e)}"}), 500

@app.route('/api/feedback-statistics', methods=['GET'])
@jwt_required()
def get_feedback_statistics():
    try:
        stats = mongodb_manager.get_feedback_statistics()
        if stats:
            return jsonify(stats)
        return jsonify({"error": "No feedback data available"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/upload-audio', methods=['POST'])
@jwt_required()
def upload_audio():
    if 'audio' not in request.files:
        return jsonify({"error": "No audio file uploaded"}), 400

    audio_file = request.files['audio']

    # 안전한 파일명 생성
    input_filename = secure_filename(audio_file.filename)
    input_filepath = os.path.join(app.config['UPLOAD_FOLDER'], input_filename)
    audio_file.save(input_filepath)

    try:
        # 오디오 파일을 WAV 16kHz mono로 변환
        converted_filepath = convert_audio_to_wav(input_filepath)

        if not converted_filepath:
            os.remove(input_filepath)
            return jsonify({"error": "오디오 파일 변환에 실패했습니다."}), 400

        # Whisper 모델을 사용하여 음성 인식 수행
        model = client.audio.transcriptions.create(
            model="whisper-1",
            file=open(converted_filepath, "rb"),
            language="ko",
            response_format="text"
        )

        # 임시 파일들 삭제
        os.remove(input_filepath)
        os.remove(converted_filepath)

        return jsonify({"text": model}), 200

    except Exception as e:
        # 임시 파일들 삭제
        if os.path.exists(input_filepath):
            os.remove(input_filepath)
        if 'converted_filepath' in locals() and os.path.exists(converted_filepath):
            os.remove(converted_filepath)

        print(f"음성 인식 중 오류 발생: {str(e)}")
        return jsonify({"error": f"음성 인식 중 오류가 발생했습니다: {str(e)}"}), 500

@app.route('/api/rate-session', methods=['POST'])
@jwt_required()
def rate_session():
    """세션에 대한 평가를 처리합니다."""
    try:
        data = request.get_json()
        session_id = data.get('session_id')
        rating = data.get('rating')
        feedback = data.get('feedback', '')

        # 필수 필드 검증
        if not all([session_id, rating]):
            return jsonify({"error": "세션 ID와 평점은 필수입니다."}), 400

        # 세션 존재 여부 확인
        session = mongodb_manager.sessions.find_one({
            'session_id': session_id,
            'user_id': request.user_id
        })

        if not session:
            return jsonify({"error": "해당 세션을 찾을 수 없습니다."}), 404

        # 이미 평가했는지 확인
        existing_rating = mongodb_manager.feedback.find_one({
            'session_id': session_id,
            'user_id': request.user_id
        })

        if existing_rating:
            return jsonify({"error": "이미 이 세션에 대해 평가하셨습니다."}), 400

        # 피드백 저장
        feedback_doc = {
            'session_id': session_id,
            'user_id': request.user_id,
            'rating': rating,
            'feedback': feedback,
            'created_at': datetime.now()
        }

        # 세션에 평가 정보 업데이트
        mongodb_manager.sessions.update_one(
            {'session_id': session_id},
            {'$set': {
                'rating': rating,
                'has_feedback': True
            }}
        )

        # 피드백 저장
        mongodb_manager.feedback.insert_one(feedback_doc)

        return jsonify({
            "message": "평가가 성공적으로 제출되었습니다.",
            "rating": rating
        }), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/generate-document', methods=['POST'])
@jwt_required()
def generate_document():
    data = request.get_json()
    transcription = data.get('transcription', '')
    summary = data.get('summary', '')

    try:
        # Create a new Document
        doc = Document()

        # Title
        title = doc.add_heading('음성 녹음 기반 소장', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Transcription Section
        doc.add_heading('원본 녹음 내용', level=1)
        transcription_para = doc.add_paragraph(transcription)
        transcription_para.style.font.size = Pt(11)

        # Summary Section
        doc.add_heading('대화 요약', level=1)
        summary_para = doc.add_paragraph(summary)
        summary_para.style.font.size = Pt(11)

        # Save document to a BytesIO object
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        return send_file(
            doc_bytes, 
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='소장.docx'
        )

    except Exception as e:
        return jsonify({
            'error': str(e),
            'status': 'error'
        }), 500

@app.route('/api/update_complaint', methods=['POST'])
@jwt_required()
def update_complaint():
    """
    사용자의 소장을 업데이트합니다.
    
    요청 본문 예시:
    {
        "complaint": "수정된 소장 내용",
        "session_id": "세션 ID (선택적)"
    }
    
    반환 예시:
    {
        "message": "소장이 성공적으로 업데이트되었습니다.",
        "updated_complaint": "수정된 소장 내용"
    }
    """
    try:
        # 요청 데이터 검증
        data = request.get_json()
        if not data or 'complaint' not in data:
            return jsonify({"error": "소장 내용이 제공되지 않았습니다."}), 400

        updated_complaint = data['complaint']
        session_id = data.get('session_id')


        # 세션 ID가 제공된 경우 해당 세션 업데이트
        if session_id:
            update_result = mongodb_manager.sessions.update_one(
                {"session_id": session_id, "user_id": request.user_id},
                {"$set": {"generated_content.complaint": updated_complaint}}
            )

            if update_result.modified_count == 0:
                # 세션을 찾지 못하거나 업데이트 실패
                return jsonify({
                    "error": "해당 세션을 찾을 수 없거나 업데이트 권한이 없습니다.",
                    "details": f"Session ID: {session_id}, User ID: {request.user_id}"
                }), 403

        # 추가적인 로깅 또는 처
        print(f"소장 업데이트 - 사용자 ID: {request.user_id}, 세션 ID: {session_id}")

        return jsonify({
            "message": "소장이 성공적으로 업데이트되었습니다.",
            "updated_complaint": updated_complaint
        }), 200

    except Exception as e:
        # 예외 처리
        print(f"소장 업데이트 중 오류 발생: {str(e)}")
        return jsonify({
            "error": "소장 업데이트 중 예기치 않은 오류가 발생했습니다.",
            "details": str(e)
        }), 500

@app.route('/api/start-consultation', methods=['POST'])
@jwt_required()
def start_consultation():
    """
    새 상담 세션을 시작하고 초기 대화 내용을 저장합니다.
    
    요청 본문 예시:
    {
        "consultation_text": "상담 내용",
        "conversation": "상담 내용 (대화 기록용)"
    }
    
    반환 예시:
    {
        "session_id": "새로 생성된 세션 ID",
        "consultation_text": "상담 내용",
        "created_at": "세션 생성 시간"
    }
    """
    try:
        # 요청 데이터 파싱
        data = request.get_json()
        consultation_text = data.get('consultation_text')
        conversation = data.get('conversation', consultation_text)

        # 입력 검증
        if not consultation_text:
            return jsonify({"error": "상담 내용이 필요합니다."}), 400

        # 세션 생성
        session_id = str(uuid.uuid4())

        # 세션 저장
        session_data = {
            "_id": session_id,
            "user_id": request.user_id,
            "consultation_text": consultation_text,
            "conversation": conversation,
            "created_at": datetime.utcnow(),
            "status": "in_progress"
        }

        # MongoDB에 세션 저장
        mongodb_manager.sessions.insert_one(session_data)

        # 응답 생성
        return jsonify({
            "session_id": session_id,
            "consultation_text": consultation_text,
            "created_at": session_data['created_at'].isoformat()
        }), 201

    except Exception as e:
        print(f"상담 시작 중 오류 발생: {str(e)}")
        return jsonify({"error": "상담 시작 중 오류가 발생했습니다.", "details": str(e)}), 500



def convert_audio_to_wav(input_file):
    """
    Convert input audio file to WAV 16kHz mono using FFmpeg
    
    Args:
        input_file (str): Path to the input audio file
    
    Returns:
        str: Path to the converted WAV file, or None if conversion fails
    """
    try:
        # Generate a unique filename for the converted file
        output_filename = f"{uuid.uuid4()}.wav"
        output_file = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)

        # FFmpeg command to convert audio to WAV 16kHz mono
        ffmpeg_command = [
            'ffmpeg', 
            '-i', input_file,  # Input file
            '-acodec', 'pcm_s16le',  # 16-bit PCM
            '-ar', '16000',  # Sample rate 16kHz
            '-ac', '1',  # Mono audio
            output_file
        ]

        # Run FFmpeg conversion
        result = subprocess.run(ffmpeg_command, capture_output=True, text=True)

        # Check if conversion was successful
        if result.returncode == 0 and os.path.exists(output_file):
            return output_file
        else:
            # Log FFmpeg error output
            print(f"FFmpeg conversion error: {result.stderr}")
            return None

    except Exception as e:
        print(f"Audio conversion error: {e}")
        return None

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
