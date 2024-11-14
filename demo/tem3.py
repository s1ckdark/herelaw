import openai
import json
from typing import Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
import os
import re

class DivorceComplaintGenerator:
    def __init__(self):
        load_dotenv()  # Load environment variables from .env file
        open_api_key = os.getenv("HERELAW_OPENAI_API_KEY")
        if not open_api_key:
            raise ValueError("API key not found. Please add it to the .env file.")
        openai.api_key = open_api_key
        self.doc = Document()
        self.client = openai.OpenAI(api_key=open_api_key)
        self.section_counter = 1
        # JSON 템플릿 파일 로드
        with open('complaint_template.json', 'r', encoding='utf-8') as f:
            self.template = json.load(f)

    def read_dialog_from_docx(self, file_path: str) -> str:
        """docx 파일에서 대화 내용 읽기"""
        doc = Document(file_path)
        dialog_text = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        return '\n'.join(dialog_text)

    def extract_information_from_dialog(self, dialog_text: str) -> Dict[str, Any]:
        """대화에서 필요한 정보를 추출하여 템플릿에 맞게 데이터를 반환"""
        
        system_message = """
        추출해야할 정보: 
        {
            "basic_info": {
                "case_name": null,
                "court_name": null,
                "date_created": null
            },
            "parties": {
                "plaintiff": {
                    "name": null,
                    "registration_number": null,
                    "domicile": null,
                    "address": null,
                    "postal_code": null
                },
                "plaintiff_representative": {
                    "type": "lawyer",
                    "name": null,
                    "office_name": null,
                    "address": null,
                    "contact": {
                        "phone": null,
                        "fax": null
                    }
                },
                "defendant": [
                    {
                        "order": 1,
                        "name": null,
                        "registration_number": null,
                        "domicile": null,
                        "address": null,
                        "postal_code": null
                    }
                ],
                "case_subject": [
                    {
                        "order": null,
                        "name": null,
                        "registration_number": null,
                        "domicile": null,
                        "address": null
                    }
                ]
            },
            "claim_purpose": {
                "divorce_claim": {
                    "claim": false
                },
                "alimony": {
                    "claim": false,
                    "amount": null,
                    "interest": {
                        "rate": null,
                        "start_point": "day_after_service",
                        "end_point": "payment_date"
                    },
                    "joint_responsibility": false
                },
                "property_division": {
                    "claim": false,
                    "amount": null,
                    "interest": {
                        "rate": null,
                        "start_point": "day_after_verdict",
                        "end_point": "payment_date"
                    }
                },
                "custody_designation": {
                    "claim": false,
                    "designated_person": null
                },
                "guardian_designation": {
                    "claim": false,
                    "designated_person": null
                },
                "child_support": {
                    "claim": false,
                    "payer": null,
                    "child_details": [
                        {
                            "child_order": null,
                            "monthly_amount": null,
                            "payment_period": {
                                "start_date": "day_after_service",
                                "end_date": null
                            }
                        }
                    ],
                    "payment_day": "end_of_month"
                },
                "litigation_cost": {
                    "payer": null
                },
                "provisional_execution": {
                    "request": false,
                    "target_items": []
                }
            },
            "claim_reason": {
                "relationship_between_parties": {
                    "marriage": {
                        "registration_date": null,
                        "marriage_duration": null,
                        "number_of_children": null
                    }
                },
                "divorce_reason": {
                    "legal_basis": {
                        "law": "Civil Code",
                        "article": "Article 840",
                        "clause": null
                    },
                    "reason_details": [
                        {
                            "type": null,
                            "detailed_reason": null,
                            "time_period": {
                                "start": null,
                                "end": null
                            }
                        }
                    ]
                },
                "alimony_claim_reason": {
                    "reason_for_claim": [],
                    "amount_basis": []
                },
                "property_division_claim_reason": {
                    "reason_for_claim": [],
                    "contribution_to_property": [],
                    "amount_basis": []
                },
                "custody_and_guardianship_reason": {
                    "reason_for_claim": [],
                    "eligibility_basis": []
                }
            },
            "evidence_methods": {
                "required_documents": {
                    "marriage_certificate": {
                        "submitted": false,
                        "document_number": null
                    },
                    "family_relation_certificate": {
                        "submitted": false,
                        "document_number": null
                    },
                    "resident_registration": {
                        "submitted": false,
                        "document_number": null
                    }
                },
                "other_evidence": [
                    {
                        "evidence_number": null,
                        "evidence_name": null,
                        "purpose": null,
                        "submitted": false
                    }
                ]
            },
            "attachments": {
                "power_of_attorney": {
                    "submitted": false
                },
                "proof_of_service_fee_payment": {
                    "submitted": false,
                    "payment_amount": null
                },
                "proof_of_stamp_attachment": {
                    "submitted": false,
                    "stamp_amount": null
                }
            }
        }

        이혼에 관한 변호사와의 대화이므로 양육권과 위자료 혹은 재산 분할에 대한 내용을 확인 부탁합니다.
        JSON 형식을 엄격히 준수하고, 결과는 JSON 데이터 형식으로만 응답해 주세요.
        항목들에 대해서는 추론할 수 있는 내용이 있다면 데이터를 추출하여 이혼 소장에 어울리는 한글 문장으로 표기해주세요.
        입증서류에는 혼인관계증명서가 필수서류로, 양육권 및 비용에 대한 내용이 있으면 가족관계증명서가 입증서류에 추가된다.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": dialog_text}
                ],
                temperature=.5,
            )
            
            # 응답을 JSON 형식으로 파싱
            content = re.sub(r"^```(?:json)?|```$", "", response.choices[0].message.content.strip(), flags=re.MULTILINE)
            filled_data = json.loads(content)
            print(filled_data)
            return filled_data
        
        except json.JSONDecodeError as e:
            print("JSON 디코딩 오류:", e)
            return None
        except openai.OpenAIError as e:
            print("OpenAI API 요청 오류:", e)
            return None

    def add_section_title(self, title, level):
        """레벨에 따라 제목 추가 (자동 번호 추가 및 줄띄움)"""
        if level == 2:
            title_text = f"{self.section_counter}. {title}"
            self.section_counter += 1
        elif level == 3:
            title_text = f"{self.section_counter - 1}.{self.subsection_counter}. {title}"
            self.subsection_counter += 1
        else:
            title_text = title

        # 제목 추가
        if level < 4:
            self.doc.add_paragraph("\n")  # 위쪽 한 줄 띄움
            self.doc.add_heading(title_text, level=level)
            self.doc.add_paragraph("\n")  # 아래쪽 한 줄 띄움
        else:
            self.doc.add_paragraph(title_text, style="Heading 4")
            self.doc.add_paragraph()  # 줄내림 추가

    def add_alpha_list_item(self, text, index, indent_level=1):
        """가, 나, 다... 순서로 하위 항목 추가 (들여쓰기 포함)"""
        alpha_list = ["가", "나", "다", "라", "마", "바", "사", "아", "자", "차"]
        paragraph = self.doc.add_paragraph(f"{alpha_list[index]}. {text}")
        paragraph.paragraph_format.left_indent = Pt(20 * indent_level)  # 들여쓰기 적용

    def add_numbered_list_item(self, text, number, indent_level=1):
        """숫자 순서로 하위 항목 추가 (들여쓰기 포함)"""
        paragraph = self.doc.add_paragraph(f"{number}. {text}")
        paragraph.paragraph_format.left_indent = Pt(20 * indent_level)  # 들여쓰기 적용

    def add_section_title(self, title, level=2):
        self.doc.add_heading(title, level=level)

    def create_divorce_complaint(self, data):
        doc = self.doc

        # 문서 제목 및 기본 정보
        doc.add_heading("이혼 청구의 소", level=1)
        doc.add_paragraph("소 장\n")

        # 기본 정보
        basic_info = data["basic_info"]
        doc.add_paragraph(f"사건명: {basic_info.get('case_name', '미기재')}")
        doc.add_paragraph(f"법원: {basic_info.get('court_name', '미기재')}")
        doc.add_paragraph(f"작성일자: {basic_info.get('date_created', datetime.now().strftime('%Y-%m-%d'))}\n")

        # 원고 정보
        plaintiff = data["parties"]["plaintiff"]
        doc.add_paragraph(f"원고: {plaintiff.get('name')} (주민등록번호: {plaintiff.get('registration_number', '미기재')})")
        doc.add_paragraph(f"등록기준지: {plaintiff.get('domicile', '미기재')}")
        doc.add_paragraph(f"주소: {plaintiff.get('address', '미기재')} (우편번호: {plaintiff.get('postal_code', '미기재')})\n")

        # 원고 소송대리인 정보
        representative = data["parties"]["plaintiff_representative"]
        doc.add_paragraph(f"원고 소송대리인: 변호사 {representative.get('name', '미기재')}")
        doc.add_paragraph(f"{representative.get('office_name', '미기재')}")
        doc.add_paragraph(f"{representative.get('address', '미기재')}")
        doc.add_paragraph(f"(전화: {representative['contact'].get('phone', '미기재')} 팩스: {representative['contact'].get('fax', '미기재')})\n")

        # 피고 정보
        for idx, defendant in enumerate(data["parties"]["defendant"], 1):
            doc.add_paragraph(f"피 고 {idx}. {defendant.get('name')} (주민등록번호: {defendant.get('registration_number', '미기재')})")
            doc.add_paragraph(f"등록기준지: {defendant.get('domicile', '미기재')}")
            doc.add_paragraph(f"주소: {defendant.get('address', '미기재')} (우편번호: {defendant.get('postal_code', '미기재')})\n")

        # 사건 본인 정보
        for idx, case_subject in enumerate(data["parties"]["case_subject"], 1):
            doc.add_paragraph(f"사 건 본 인 {idx}. {case_subject.get('name', '미기재')} (주민등록번호: {case_subject.get('registration_number', '미기재')})")
            doc.add_paragraph(f"등록기준지: {case_subject.get('domicile', '미기재')}")
            doc.add_paragraph(f"주소: {case_subject.get('address', '미기재')}\n")

        # 청구 취지
        self.add_section_title("청 구 취 지", level=2)
        claim_purpose = data["claim_purpose"]
        numbered_item_index = 1

        if claim_purpose["divorce_claim"].get("claim"):
            self.add_numbered_list_item("원고와 피고는 이혼한다.", numbered_item_index)
            numbered_item_index += 1

        alimony = claim_purpose["alimony"]
        if alimony.get("claim"):
            alimony_text = (
                f"피고들은 연대하여 원고에게 위자료로 금 {alimony.get('amount')}원 및 이에 대하여 "
                f"소장 부본 송달일 다음 날부터 다 갚는 날까지 연 {alimony['interest'].get('rate')}%의 "
                f"비율로 계산한 돈을 지급하라."
            )
            self.add_numbered_list_item(alimony_text, numbered_item_index)
            numbered_item_index += 1

        property_division = claim_purpose["property_division"]
        if property_division.get("claim"):
            property_division_text = (
                f"피고는 원고에게 재산분할로 금 {property_division.get('amount')}원 및 이에 대한 판결 확정일 "
                f"다음 날부터 다 갚는 날까지 연 {property_division['interest'].get('rate')}%의 비율로 계산한 돈을 지급하라."
            )
            self.add_numbered_list_item(property_division_text, numbered_item_index)
            numbered_item_index += 1

        if claim_purpose["custody_designation"].get("claim"):
            custody_text = f"사건본인들의 친권행사자 및 양육자로 {claim_purpose['custody_designation'].get('designated_person', '미기재')}을 지정한다."
            self.add_numbered_list_item(custody_text, numbered_item_index)
            numbered_item_index += 1

        child_support = claim_purpose["child_support"]
        if child_support.get("claim"):
            child_support_text = "피고는 사건본인들의 양육비로 매월 다음과 같이 지급하라:"
            self.add_numbered_list_item(child_support_text, numbered_item_index)
            for child in child_support["child_details"]:
                child_order = child.get("child_order", "미기재")
                monthly_amount = child.get("monthly_amount", "미기재")
                doc.add_paragraph(f" - 사건본인 {child_order}: 월 {monthly_amount}원")
            numbered_item_index += 1

        # 소송비용 부담 및 가집행 청구 항목 추가
        if claim_purpose["litigation_cost"].get("payer"):
            self.add_numbered_list_item("소송비용은 피고가 부담한다.", numbered_item_index)
            numbered_item_index += 1

        if claim_purpose["provisional_execution"].get("request"):
            self.add_numbered_list_item("위 제2항 및 제3항은 가집행할 수 있다.", numbered_item_index)

        # 청구 원인
        self.add_section_title("청 구 원 인", level=2)
        self.add_section_title("당사자의 지위", level=3)
        marriage_info = data["claim_reason"]["relationship_between_parties"]["marriage"]
        doc.add_paragraph(f"원고와 피고는 {marriage_info.get('registration_date', '미기재')} 혼인신고를 마친 법률상 부부로, 혼인 기간은 {marriage_info.get('marriage_duration', '미기재')}이며 자녀는 {marriage_info.get('number_of_children', '미기재')}명입니다.\n")

        # 재판상 이혼 사유
        self.add_section_title("재판상 이혼 사유", level=3)
        divorce_reason = data["claim_reason"]["divorce_reason"]["reason_details"]
        for i, reason in enumerate(divorce_reason):
            self.add_alpha_list_item(f"{reason.get('type')} - {reason.get('detailed_reason')}", i, indent_level=2)

        # 결론 및 하위 항목
        self.add_section_title("결 론", level=2)
        doc.add_paragraph("위와 같은 이유로 원고는 청구취지 기재와 같은 판결을 구하고자 이 건 소를 제기합니다.")

        # 입증 방법
        self.add_section_title("입 증 방 법", level=3)
        evidence = data["evidence_methods"]["other_evidence"][0]
        doc.add_paragraph(f"갑 제1호증: {evidence.get('evidence_name')}")

        # 첨부 서류
        self.add_section_title("첨 부 서 류", level=3)
        attachments = data["attachments"]
        if attachments["power_of_attorney"]["submitted"]:
            doc.add_paragraph("소송위임장: 제출")
        if attachments["proof_of_service_fee_payment"]["submitted"]:
            doc.add_paragraph("송달료 납부 증명서: 제출")
        if attachments["proof_of_stamp_attachment"]["submitted"]:
            doc.add_paragraph("인지 첨부 증명서: 제출")

        # 작성 날짜 및 원고 소송대리인
        doc.add_paragraph("\n작성일자: " + datetime.now().strftime('%Y년 %m월 %d일'))
        doc.add_paragraph("원고 소송대리인: " + representative.get('name'))

        return doc
    
def main():
    generator = DivorceComplaintGenerator()

    dialog_path = "data.docx"  # data.docx 파일 경로
    output_path = "result_temp.docx"  # result.docx 파일 경로

    dialog_text = generator.read_dialog_from_docx(dialog_path)
    extracted_data = generator.extract_information_from_dialog(dialog_text)

    doc = generator.create_divorce_complaint(extracted_data)
    for paragraph in doc.paragraphs:
        print(paragraph.text)

if __name__ == "__main__":
    main()
